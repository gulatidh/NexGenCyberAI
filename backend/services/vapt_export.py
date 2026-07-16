"""
VAPT Report Export Service
Generates professional PDF and DOCX reports for VAPT engagements.
"""
import io
import json
from datetime import datetime
from typing import Any, Dict, List, Optional

# ── ReportLab imports ──────────────────────────────────────────────────────────
from reportlab.lib import colors
from reportlab.lib.colors import HexColor
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm, mm
from reportlab.platypus import (
    BaseDocTemplate,
    Frame,
    HRFlowable,
    Image,
    KeepTogether,
    PageBreak,
    PageTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
)
from reportlab.platypus.flowables import Flowable

# ── python-docx imports ────────────────────────────────────────────────────────
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.table import WD_ALIGN_VERTICAL

# ── Colour palette ─────────────────────────────────────────────────────────────
NAVY = HexColor("#1A237E")
BLUE = HexColor("#1565C0")
GREY_BG = HexColor("#ECEFF1")
DARK_TEXT = HexColor("#37474F")
WHITE = colors.white

SEV_COLORS = {
    "critical":      HexColor("#C62828"),
    "high":          HexColor("#E64A19"),
    "medium":        HexColor("#F9A825"),
    "low":           HexColor("#2E7D32"),
    "informational": HexColor("#1565C0"),
    "info":          HexColor("#1565C0"),
}

SEV_TEXT_COLORS = {
    "critical":      WHITE,
    "high":          WHITE,
    "medium":        HexColor("#37474F"),  # dark text on yellow
    "low":           WHITE,
    "informational": WHITE,
    "info":          WHITE,
}

RETEST_COLORS = {
    "passed":         HexColor("#2E7D32"),
    "failed":         HexColor("#C62828"),
    "pending":        HexColor("#757575"),
    "not_applicable": HexColor("#455A64"),
}

SEV_ORDER = ["critical", "high", "medium", "low", "informational", "info"]


def _sev_key(f: Dict) -> int:
    s = (f.get("severity") or "").lower()
    try:
        return SEV_ORDER.index(s)
    except ValueError:
        return 99


def _fmt_date(val: Any) -> str:
    if not val:
        return "—"
    if isinstance(val, datetime):
        return val.strftime("%d %B %Y")
    try:
        return datetime.fromisoformat(str(val)).strftime("%d %B %Y")
    except Exception:
        return str(val)


def _safe(val: Any, fallback: str = "—") -> str:
    if val is None or (isinstance(val, str) and not val.strip()):
        return fallback
    return str(val)


def _sev_counts(findings: List[Dict]) -> Dict[str, int]:
    counts = {s: 0 for s in ["critical", "high", "medium", "low", "informational"]}
    for f in findings:
        s = (f.get("severity") or "").lower()
        if s == "info":
            s = "informational"
        if s in counts:
            counts[s] += 1
    return counts


def _retest_counts(findings: List[Dict]) -> Dict[str, int]:
    counts = {s: 0 for s in ["passed", "failed", "pending", "not_applicable"]}
    for f in findings:
        rs = (f.get("retest_status") or "pending").lower()
        if rs in counts:
            counts[rs] += 1
    return counts


# ══════════════════════════════════════════════════════════════════════════════
# PDF GENERATION
# ══════════════════════════════════════════════════════════════════════════════

def _build_pdf_styles():
    styles = getSampleStyleSheet()
    base = dict(fontName="Helvetica", fontSize=10, leading=14, textColor=DARK_TEXT)

    def _ps(name, **kw):
        merged = {**base, **kw}
        return ParagraphStyle(name, **merged)

    return {
        "normal":       _ps("vapt_normal"),
        "bold":         _ps("vapt_bold", fontName="Helvetica-Bold"),
        "small":        _ps("vapt_small", fontSize=8, leading=10),
        "small_bold":   _ps("vapt_small_bold", fontSize=8, leading=10, fontName="Helvetica-Bold"),
        "section":      _ps("vapt_section", fontName="Helvetica-Bold", fontSize=13, leading=16, textColor=WHITE),
        "subsection":   _ps("vapt_subsection", fontName="Helvetica-Bold", fontSize=11, leading=14, textColor=BLUE),
        "finding_hdr":  _ps("vapt_finding_hdr", fontName="Helvetica-Bold", fontSize=12, leading=15, textColor=WHITE),
        "cover_title":  _ps("vapt_cover_title", fontName="Helvetica-Bold", fontSize=28, leading=34, textColor=WHITE, alignment=TA_CENTER),
        "cover_sub":    _ps("vapt_cover_sub", fontName="Helvetica", fontSize=13, leading=17, textColor=HexColor("#B0BEC5"), alignment=TA_CENTER),
        "center":       _ps("vapt_center", alignment=TA_CENTER),
        "center_bold":  _ps("vapt_center_bold", fontName="Helvetica-Bold", alignment=TA_CENTER),
        "right":        _ps("vapt_right", alignment=TA_RIGHT),
        "label":        _ps("vapt_label", fontName="Helvetica-Bold", fontSize=9, textColor=HexColor("#546E7A")),
        "value":        _ps("vapt_value", fontSize=10),
        "bullet":       _ps("vapt_bullet", fontSize=10, leftIndent=12, spaceBefore=2),
        "code":         _ps("vapt_code", fontName="Courier", fontSize=9, leading=12, backColor=HexColor("#F5F5F5"), leftIndent=8, rightIndent=8),
    }


class _PageNumCanvas:
    """Mixin helper — not used directly; canvas callbacks handle numbering."""
    pass


def generate_pdf(report: Dict, findings: List[Dict], client_name: str) -> bytes:
    buf = io.BytesIO()
    styles = _build_pdf_styles()
    sorted_findings = sorted(findings, key=_sev_key)
    sev_counts = _sev_counts(findings)
    retest_counts = _retest_counts(findings)

    # Scope / methodology JSON
    scope = {}
    methodology = {}
    try:
        scope = json.loads(report.get("scope_json") or "{}")
    except Exception:
        pass
    try:
        methodology = json.loads(report.get("methodology_json") or "{}")
    except Exception:
        pass

    # ── Page templates (cover vs content) ────────────────────────────────────
    PAGE_W, PAGE_H = A4
    MARGIN = 2 * cm

    report_title = _safe(report.get("title"), "VAPT Report")
    classification = _safe(report.get("classification"), "Confidential")
    version_str = _safe(report.get("version"), "1.0")

    def _cover_template(canvas, doc):
        canvas.saveState()
        # Dark navy background
        canvas.setFillColor(HexColor("#0D1B4B"))
        canvas.rect(0, 0, PAGE_W, PAGE_H, fill=1, stroke=0)
        # Top accent bar
        canvas.setFillColor(HexColor("#1565C0"))
        canvas.rect(0, PAGE_H - 0.8 * cm, PAGE_W, 0.8 * cm, fill=1, stroke=0)
        # Bottom bar
        canvas.setFillColor(HexColor("#1565C0"))
        canvas.rect(0, 0, PAGE_W, 0.8 * cm, fill=1, stroke=0)
        # Left accent strip
        canvas.setFillColor(HexColor("#C62828"))
        canvas.rect(0, 0.8 * cm, 0.5 * cm, PAGE_H - 1.6 * cm, fill=1, stroke=0)
        canvas.restoreState()

    def _content_template(canvas, doc):
        canvas.saveState()
        # Header line
        canvas.setFillColor(NAVY)
        canvas.rect(MARGIN, PAGE_H - MARGIN + 0.2 * cm, PAGE_W - 2 * MARGIN, 0.05 * cm, fill=1, stroke=0)
        # Header text left
        canvas.setFont("Helvetica", 8)
        canvas.setFillColor(DARK_TEXT)
        canvas.drawString(MARGIN, PAGE_H - MARGIN + 0.45 * cm, report_title[:80])
        # Header text right
        hdr_right = f"{classification}  |  Page {doc.page}"
        canvas.drawRightString(PAGE_W - MARGIN, PAGE_H - MARGIN + 0.45 * cm, hdr_right)
        # Footer line
        canvas.setFillColor(NAVY)
        canvas.rect(MARGIN, MARGIN - 0.4 * cm, PAGE_W - 2 * MARGIN, 0.05 * cm, fill=1, stroke=0)
        # Footer text
        canvas.setFont("Helvetica", 8)
        canvas.setFillColor(DARK_TEXT)
        canvas.drawString(MARGIN, MARGIN - 0.6 * cm, client_name)
        canvas.drawRightString(PAGE_W - MARGIN, MARGIN - 0.6 * cm, f"Version {version_str}")
        canvas.restoreState()

    cover_frame = Frame(0, 0, PAGE_W, PAGE_H, leftPadding=3 * cm, rightPadding=1.5 * cm,
                        topPadding=2 * cm, bottomPadding=2 * cm)
    content_frame = Frame(MARGIN, MARGIN, PAGE_W - 2 * MARGIN, PAGE_H - 2 * MARGIN,
                          topPadding=0.8 * cm, bottomPadding=0.8 * cm)

    doc = BaseDocTemplate(
        buf, pagesize=A4,
        leftMargin=MARGIN, rightMargin=MARGIN,
        topMargin=MARGIN, bottomMargin=MARGIN,
    )
    doc.addPageTemplates([
        PageTemplate(id="Cover", frames=[cover_frame], onPage=_cover_template),
        PageTemplate(id="Content", frames=[content_frame], onPage=_content_template),
    ])

    story = []

    # ── Helper functions ──────────────────────────────────────────────────────
    def section_header(text: str) -> Table:
        """Full-width navy section header table."""
        tbl = Table([[Paragraph(text, styles["section"])]], colWidths=[PAGE_W - 2 * MARGIN])
        tbl.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, -1), NAVY),
            ("LEFTPADDING", (0, 0), (-1, -1), 10),
            ("RIGHTPADDING", (0, 0), (-1, -1), 10),
            ("TOPPADDING", (0, 0), (-1, -1), 6),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
        ]))
        return tbl

    def subsection_header(text: str) -> Paragraph:
        return Paragraph(text, styles["subsection"])

    def labeled_row(label: str, value: str):
        return [Paragraph(label, styles["label"]), Paragraph(_safe(value), styles["normal"])]

    def sev_cell(sev: str) -> Table:
        s = sev.lower()
        bg = SEV_COLORS.get(s, HexColor("#757575"))
        tc = SEV_TEXT_COLORS.get(s, WHITE)
        label = sev.capitalize()
        p = ParagraphStyle("sev_cell", fontName="Helvetica-Bold", fontSize=9,
                           textColor=tc, alignment=TA_CENTER)
        t = Table([[Paragraph(label, p)]], colWidths=[2.5 * cm])
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, -1), bg),
            ("TOPPADDING", (0, 0), (-1, -1), 3),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ("LEFTPADDING", (0, 0), (-1, -1), 4),
            ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ]))
        return t

    # ═══════════════════════════════════════════════════════════════════════
    # COVER PAGE
    # ═══════════════════════════════════════════════════════════════════════
    story.append(Spacer(1, 4 * cm))
    story.append(Paragraph("MONITARA AI", ParagraphStyle(
        "brand", fontName="Helvetica-Bold", fontSize=16, textColor=HexColor("#42A5F5"),
        alignment=TA_CENTER, letterSpacing=4)))
    story.append(Spacer(1, 0.5 * cm))
    story.append(Paragraph("Cybersecurity Platform", ParagraphStyle(
        "brand_sub", fontName="Helvetica", fontSize=11, textColor=HexColor("#90CAF9"),
        alignment=TA_CENTER)))
    story.append(Spacer(1, 2 * cm))
    story.append(HRFlowable(width="80%", thickness=1, color=HexColor("#1565C0"),
                            hAlign="CENTER"))
    story.append(Spacer(1, 1.5 * cm))
    story.append(Paragraph(report_title.upper(), styles["cover_title"]))
    story.append(Spacer(1, 0.8 * cm))
    story.append(Paragraph("Vulnerability Assessment &amp; Penetration Testing Report",
                            styles["cover_sub"]))
    story.append(Spacer(1, 2 * cm))

    cover_data = [
        ["Client", _safe(client_name)],
        ["Classification", classification],
        ["Version", f"v{version_str}"],
        ["Report Date", _fmt_date(report.get("report_date"))],
        ["Prepared By", _safe(report.get("prepared_by"))],
        ["Reviewed By", _safe(report.get("reviewed_by"))],
        ["Status", _safe(report.get("status"), "Draft").capitalize()],
    ]
    cover_style_normal = ParagraphStyle("c_n", fontName="Helvetica", fontSize=10, textColor=WHITE)
    cover_style_bold = ParagraphStyle("c_b", fontName="Helvetica-Bold", fontSize=10, textColor=HexColor("#90CAF9"))
    cover_tbl_data = [[Paragraph(r[0], cover_style_bold), Paragraph(r[1], cover_style_normal)]
                      for r in cover_data]
    cover_tbl = Table(cover_tbl_data, colWidths=[4 * cm, 10 * cm])
    cover_tbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), HexColor("#0A1340")),
        ("LINEBELOW", (0, 0), (-1, -2), 0.5, HexColor("#1E3A8A")),
        ("TOPPADDING", (0, 0), (-1, -1), 7),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
        ("LEFTPADDING", (0, 0), (-1, -1), 10),
        ("RIGHTPADDING", (0, 0), (-1, -1), 10),
    ]))
    story.append(cover_tbl)
    story.append(Spacer(1, 2 * cm))
    story.append(Paragraph(
        "CONFIDENTIAL — This document contains sensitive security information. "
        "Distribution is restricted to authorised personnel only.",
        ParagraphStyle("conf", fontName="Helvetica-Oblique", fontSize=9,
                       textColor=HexColor("#90A4AE"), alignment=TA_CENTER)))
    story.append(PageBreak())

    # Switch to content template
    from reportlab.platypus import NextPageTemplate
    story.append(NextPageTemplate("Content"))

    # ═══════════════════════════════════════════════════════════════════════
    # SECTION 1: EXECUTIVE SUMMARY
    # ═══════════════════════════════════════════════════════════════════════
    story.append(section_header("1. Executive Summary"))
    story.append(Spacer(1, 0.4 * cm))

    # Severity statistics table
    sev_header_style = ParagraphStyle("sh", fontName="Helvetica-Bold", fontSize=11, textColor=WHITE, alignment=TA_CENTER)
    sev_count_style = ParagraphStyle("sc", fontName="Helvetica-Bold", fontSize=22, textColor=WHITE, alignment=TA_CENTER)
    sev_label_style = ParagraphStyle("sl", fontName="Helvetica", fontSize=9, textColor=WHITE, alignment=TA_CENTER)

    sev_display = [
        ("CRITICAL", "critical", sev_counts["critical"]),
        ("HIGH", "high", sev_counts["high"]),
        ("MEDIUM", "medium", sev_counts["medium"]),
        ("LOW", "low", sev_counts["low"]),
        ("INFO", "informational", sev_counts["informational"]),
        ("TOTAL", None, sum(sev_counts.values())),
    ]

    sev_row = []
    for label, key, count in sev_display:
        bg = SEV_COLORS.get(key, NAVY) if key else NAVY
        cell = Table([
            [Paragraph(label, sev_label_style)],
            [Paragraph(str(count), sev_count_style)],
        ], colWidths=[2.5 * cm])
        cell.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, -1), bg),
            ("TOPPADDING", (0, 0), (-1, -1), 8),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
            ("LEFTPADDING", (0, 0), (-1, -1), 4),
            ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ]))
        sev_row.append(cell)

    sev_tbl = Table([sev_row], colWidths=[2.5 * cm] * 6)
    sev_tbl.setStyle(TableStyle([("ALIGN", (0, 0), (-1, -1), "CENTER")]))
    story.append(sev_tbl)
    story.append(Spacer(1, 0.4 * cm))

    # Retest summary if any retests exist
    total_retest = sum(retest_counts.values())
    if total_retest > 0 and (retest_counts["passed"] + retest_counts["failed"]) > 0:
        retest_row_data = [
            [Paragraph("PASSED", sev_label_style), Paragraph(str(retest_counts["passed"]), sev_count_style)],
            [Paragraph("FAILED", sev_label_style), Paragraph(str(retest_counts["failed"]), sev_count_style)],
            [Paragraph("PENDING", sev_label_style), Paragraph(str(retest_counts["pending"]), sev_count_style)],
            [Paragraph("N/A", sev_label_style), Paragraph(str(retest_counts["not_applicable"]), sev_count_style)],
        ]
        retest_row = []
        for (lbl_p, cnt_p), (key) in zip(retest_row_data, ["passed", "failed", "pending", "not_applicable"]):
            bg = RETEST_COLORS.get(key, HexColor("#757575"))
            cell = Table([[lbl_p], [cnt_p]], colWidths=[2.5 * cm])
            cell.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, -1), bg),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
            ]))
            retest_row.append(cell)
        story.append(subsection_header("Retest Summary"))
        story.append(Spacer(1, 0.2 * cm))
        story.append(Table([retest_row], colWidths=[2.5 * cm] * 4))
        story.append(Spacer(1, 0.4 * cm))

    story.append(Spacer(1, 0.3 * cm))
    exec_summary = _safe(report.get("executive_summary"),
                         "No executive summary provided.")
    for para in exec_summary.split("\n"):
        if para.strip():
            story.append(Paragraph(para.strip(), styles["normal"]))
            story.append(Spacer(1, 0.15 * cm))

    story.append(Spacer(1, 0.5 * cm))

    # ═══════════════════════════════════════════════════════════════════════
    # SECTION 2: SCOPE
    # ═══════════════════════════════════════════════════════════════════════
    story.append(section_header("2. Scope of Engagement"))
    story.append(Spacer(1, 0.4 * cm))

    scope_meta = [
        ["Engagement Type", _safe(scope.get("engagement_type"))],
        ["Period Start", _safe(scope.get("period_start"))],
        ["Period End", _safe(scope.get("period_end"))],
    ]
    scope_tbl = Table([[Paragraph(r[0], styles["label"]), Paragraph(r[1], styles["normal"])]
                       for r in scope_meta], colWidths=[4 * cm, 13 * cm])
    scope_tbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (0, -1), GREY_BG),
        ("LINEBELOW", (0, 0), (-1, -1), 0.5, HexColor("#CFD8DC")),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
    ]))
    story.append(scope_tbl)
    story.append(Spacer(1, 0.4 * cm))

    in_scope = scope.get("in_scope") or []
    out_scope = scope.get("out_of_scope") or []

    if in_scope:
        story.append(subsection_header("In-Scope Assets"))
        story.append(Spacer(1, 0.2 * cm))
        in_scope_data = [[Paragraph(str(i + 1), styles["small_bold"]),
                          Paragraph(_safe(item), styles["normal"])]
                         for i, item in enumerate(in_scope)]
        in_scope_tbl = Table(in_scope_data, colWidths=[1 * cm, 16 * cm])
        in_scope_tbl.setStyle(TableStyle([
            ("ROWBACKGROUNDS", (0, 0), (-1, -1), [WHITE, GREY_BG]),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ("LEFTPADDING", (0, 0), (-1, -1), 8),
            ("RIGHTPADDING", (0, 0), (-1, -1), 8),
            ("LINEBELOW", (0, 0), (-1, -1), 0.3, HexColor("#CFD8DC")),
        ]))
        story.append(in_scope_tbl)
        story.append(Spacer(1, 0.3 * cm))

    if out_scope:
        story.append(subsection_header("Out-of-Scope Items"))
        story.append(Spacer(1, 0.2 * cm))
        for item in out_scope:
            story.append(Paragraph(f"• {_safe(item)}", styles["bullet"]))
        story.append(Spacer(1, 0.3 * cm))

    story.append(Spacer(1, 0.3 * cm))

    # ═══════════════════════════════════════════════════════════════════════
    # SECTION 3: METHODOLOGY
    # ═══════════════════════════════════════════════════════════════════════
    story.append(section_header("3. Methodology"))
    story.append(Spacer(1, 0.4 * cm))

    phases = methodology.get("phases") or []
    if phases:
        story.append(subsection_header("Testing Phases"))
        story.append(Spacer(1, 0.2 * cm))
        phases_header = [Paragraph(h, ParagraphStyle("ph", fontName="Helvetica-Bold",
                                                      fontSize=9, textColor=WHITE))
                         for h in ["#", "Phase", "Description"]]
        phases_data = [phases_header]
        for i, phase in enumerate(phases):
            phases_data.append([
                Paragraph(str(i + 1), styles["small_bold"]),
                Paragraph(_safe(phase.get("name")), styles["bold"]),
                Paragraph(_safe(phase.get("description")), styles["normal"]),
            ])
        phases_tbl = Table(phases_data, colWidths=[0.8 * cm, 4 * cm, 12.2 * cm])
        phases_tbl.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), NAVY),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [WHITE, GREY_BG]),
            ("LINEBELOW", (0, 0), (-1, -1), 0.3, HexColor("#CFD8DC")),
            ("TOPPADDING", (0, 0), (-1, -1), 5),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ("LEFTPADDING", (0, 0), (-1, -1), 8),
            ("RIGHTPADDING", (0, 0), (-1, -1), 8),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ]))
        story.append(phases_tbl)
        story.append(Spacer(1, 0.4 * cm))

    tools = methodology.get("tools") or []
    if tools:
        story.append(subsection_header("Tools Used"))
        story.append(Spacer(1, 0.1 * cm))
        story.append(Paragraph(" | ".join(tools), styles["normal"]))
        story.append(Spacer(1, 0.3 * cm))

    standards = methodology.get("standards") or []
    if standards:
        story.append(subsection_header("Standards &amp; Frameworks Referenced"))
        story.append(Spacer(1, 0.1 * cm))
        for std in standards:
            story.append(Paragraph(f"• {_safe(std)}", styles["bullet"]))
        story.append(Spacer(1, 0.3 * cm))

    story.append(Spacer(1, 0.3 * cm))

    # ═══════════════════════════════════════════════════════════════════════
    # SECTION 4: SEVERITY RATING MATRIX
    # ═══════════════════════════════════════════════════════════════════════
    story.append(section_header("4. Severity Rating Matrix"))
    story.append(Spacer(1, 0.4 * cm))

    matrix_hdr_style = ParagraphStyle("mhdr", fontName="Helvetica-Bold", fontSize=9, textColor=WHITE)
    matrix_body_style = styles["small"]
    matrix_hdr = [Paragraph(h, matrix_hdr_style) for h in
                  ["Rating", "CVSS Range", "Description", "Examples", "Response SLA"]]
    matrix_rows = [
        ("CRITICAL", "9.0 – 10.0", HexColor("#C62828"), WHITE,
         "Remote unauthenticated code execution, complete system compromise.",
         "RCE, SQLi (full DB access), Privilege escalation to domain admin",
         "Immediate — within 24 hours"),
        ("HIGH", "7.0 – 8.9", HexColor("#E64A19"), WHITE,
         "Significant risk requiring urgent attention.",
         "SSRF, Stored XSS, Authentication bypass, Exposed credentials",
         "Urgent — within 72 hours"),
        ("MEDIUM", "4.0 – 6.9", HexColor("#F9A825"), DARK_TEXT,
         "Moderate risk requiring planned remediation.",
         "Reflected XSS, CSRF, Insecure configuration, Sensitive data in logs",
         "Planned — within 30 days"),
        ("LOW", "0.1 – 3.9", HexColor("#2E7D32"), WHITE,
         "Minor risk with limited exposure.",
         "Missing headers, Verbose error messages, Weak cipher suites",
         "Scheduled — within 90 days"),
        ("INFORMATIONAL", "N/A", HexColor("#1565C0"), WHITE,
         "Best practice observations, no immediate risk.",
         "Documentation gaps, Minor misconfigurations, Improvement areas",
         "Best effort"),
    ]
    matrix_data = [matrix_hdr]
    for row in matrix_rows:
        label, cvss, bg, tc, desc, examples, sla = row
        row_style = ParagraphStyle("mr", fontName="Helvetica", fontSize=8, leading=11, textColor=tc)
        lbl_style = ParagraphStyle("ml", fontName="Helvetica-Bold", fontSize=9, textColor=tc)
        matrix_data.append([
            Paragraph(label, lbl_style),
            Paragraph(cvss, row_style),
            Paragraph(desc, row_style),
            Paragraph(examples, row_style),
            Paragraph(sla, row_style),
        ])

    matrix_tbl = Table(matrix_data, colWidths=[2.5 * cm, 1.8 * cm, 4 * cm, 5 * cm, 3.7 * cm])

    matrix_ts = [
        ("BACKGROUND", (0, 0), (-1, 0), NAVY),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LINEBELOW", (0, 0), (-1, -1), 0.3, HexColor("#37474F")),
    ]
    sev_names = ["critical", "high", "medium", "low", "informational"]
    for i, sev in enumerate(sev_names):
        bg = SEV_COLORS[sev]
        matrix_ts.append(("BACKGROUND", (0, i + 1), (0, i + 1), bg))

    matrix_tbl.setStyle(TableStyle(matrix_ts))
    story.append(matrix_tbl)
    story.append(Spacer(1, 0.5 * cm))

    # ═══════════════════════════════════════════════════════════════════════
    # SECTION 5: FINDINGS SUMMARY
    # ═══════════════════════════════════════════════════════════════════════
    story.append(section_header("5. Findings Summary"))
    story.append(Spacer(1, 0.4 * cm))

    if not sorted_findings:
        story.append(Paragraph("No findings recorded in this report.", styles["normal"]))
    else:
        fs_hdr_style = ParagraphStyle("fsh", fontName="Helvetica-Bold", fontSize=9, textColor=WHITE)
        fs_hdr = [Paragraph(h, fs_hdr_style) for h in
                  ["ID", "Title", "Severity", "Affected Asset", "Retest Status"]]
        fs_data = [fs_hdr]
        for f in sorted_findings:
            sev = (f.get("severity") or "").lower()
            rs = (f.get("retest_status") or "pending").lower()
            rs_bg = RETEST_COLORS.get(rs, HexColor("#757575"))
            rs_style = ParagraphStyle("rss", fontName="Helvetica-Bold", fontSize=8,
                                      textColor=WHITE, alignment=TA_CENTER)
            rs_cell = Table([[Paragraph(rs.replace("_", " ").title(), rs_style)]],
                            colWidths=[2.8 * cm])
            rs_cell.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, -1), rs_bg),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
            ]))
            fs_data.append([
                Paragraph(_safe(f.get("finding_id")), styles["small_bold"]),
                Paragraph(_safe(f.get("title")), styles["small"]),
                sev_cell(sev),
                Paragraph(_safe(f.get("affected_asset")), styles["small"]),
                rs_cell,
            ])
        fs_tbl = Table(fs_data, colWidths=[1.5 * cm, 6 * cm, 2.5 * cm, 4.2 * cm, 2.8 * cm])
        fs_tbl.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), NAVY),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [WHITE, GREY_BG]),
            ("TOPPADDING", (0, 0), (-1, -1), 5),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ("LEFTPADDING", (0, 0), (-1, -1), 6),
            ("RIGHTPADDING", (0, 0), (-1, -1), 6),
            ("LINEBELOW", (0, 0), (-1, -1), 0.3, HexColor("#CFD8DC")),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ]))
        story.append(fs_tbl)

    story.append(Spacer(1, 0.5 * cm))

    # ═══════════════════════════════════════════════════════════════════════
    # SECTION 6: DETAILED FINDINGS
    # ═══════════════════════════════════════════════════════════════════════
    story.append(section_header("6. Detailed Findings"))
    story.append(Spacer(1, 0.4 * cm))

    if not sorted_findings:
        story.append(Paragraph("No findings recorded.", styles["normal"]))
    else:
        for fi, f in enumerate(sorted_findings):
            sev = (f.get("severity") or "").lower()
            bg = SEV_COLORS.get(sev, HexColor("#757575"))
            tc = SEV_TEXT_COLORS.get(sev, WHITE)
            fid = _safe(f.get("finding_id"), f"F-{fi + 1:02d}")
            ftitle = _safe(f.get("title"))
            rs = (f.get("retest_status") or "pending").lower()

            finding_elements = []

            # Color-coded header
            hdr_style = ParagraphStyle("fdh", fontName="Helvetica-Bold", fontSize=12,
                                       textColor=tc, leading=15)
            fhdr_tbl = Table([[Paragraph(f"{fid}  —  {ftitle}", hdr_style)]],
                              colWidths=[PAGE_W - 2 * MARGIN])
            fhdr_tbl.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, -1), bg),
                ("TOPPADDING", (0, 0), (-1, -1), 8),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
                ("LEFTPADDING", (0, 0), (-1, -1), 10),
                ("RIGHTPADDING", (0, 0), (-1, -1), 10),
            ]))
            finding_elements.append(fhdr_tbl)

            # Meta row
            meta_bold = ParagraphStyle("mb", fontName="Helvetica-Bold", fontSize=9, textColor=DARK_TEXT)
            meta_val = ParagraphStyle("mv", fontName="Helvetica", fontSize=9, textColor=DARK_TEXT)
            rs_display = rs.replace("_", " ").title()
            meta_tbl = Table([
                [Paragraph("Severity:", meta_bold), Paragraph(sev.capitalize(), meta_val),
                 Paragraph("Affected Asset:", meta_bold), Paragraph(_safe(f.get("affected_asset")), meta_val),
                 Paragraph("Retest Status:", meta_bold), Paragraph(rs_display, meta_val)],
            ], colWidths=[2.2 * cm, 2.8 * cm, 3 * cm, 4 * cm, 2.5 * cm, 2.5 * cm])
            meta_tbl.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, -1), GREY_BG),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("LINEBELOW", (0, 0), (-1, -1), 0.3, HexColor("#CFD8DC")),
            ]))
            finding_elements.append(meta_tbl)
            finding_elements.append(Spacer(1, 0.3 * cm))

            def field_block(label: str, value: str, code_style: bool = False):
                if not value or value == "—":
                    return
                finding_elements.append(Paragraph(label, styles["subsection"]))
                finding_elements.append(Spacer(1, 0.1 * cm))
                st = styles["code"] if code_style else styles["normal"]
                for line in value.split("\n"):
                    if line.strip():
                        finding_elements.append(Paragraph(line, st))
                        finding_elements.append(Spacer(1, 0.05 * cm))
                finding_elements.append(Spacer(1, 0.2 * cm))

            field_block("Description", _safe(f.get("description")))
            field_block("Business Impact", _safe(f.get("impact")))
            field_block("Evidence", _safe(f.get("evidence")))
            field_block("Reproduction Steps", _safe(f.get("reproduction_steps")), code_style=True)
            field_block("Recommendation", _safe(f.get("recommendation")))
            field_block("References", _safe(f.get("references")))

            # Retest notes box
            rn = f.get("retest_notes")
            if rn and rn.strip():
                rn_bg = RETEST_COLORS.get(rs, HexColor("#757575"))
                rn_style = ParagraphStyle("rn", fontName="Helvetica", fontSize=9,
                                          textColor=WHITE, leading=13)
                rn_lbl_style = ParagraphStyle("rnl", fontName="Helvetica-Bold", fontSize=9,
                                              textColor=WHITE)
                rn_tbl = Table([
                    [Paragraph(f"Retest Notes ({rs_display})", rn_lbl_style)],
                    [Paragraph(rn, rn_style)],
                ], colWidths=[PAGE_W - 2 * MARGIN])
                rn_tbl.setStyle(TableStyle([
                    ("BACKGROUND", (0, 0), (-1, -1), rn_bg),
                    ("TOPPADDING", (0, 0), (-1, -1), 6),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                    ("LEFTPADDING", (0, 0), (-1, -1), 10),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 10),
                ]))
                finding_elements.append(rn_tbl)
                finding_elements.append(Spacer(1, 0.2 * cm))

            if fi < len(sorted_findings) - 1:
                finding_elements.append(HRFlowable(width="100%", thickness=1,
                                                   color=HexColor("#B0BEC5")))
                finding_elements.append(Spacer(1, 0.4 * cm))

            try:
                story.append(KeepTogether(finding_elements[:8]))
                for el in finding_elements[8:]:
                    story.append(el)
            except Exception:
                for el in finding_elements:
                    story.append(el)

    story.append(Spacer(1, 0.5 * cm))

    # ═══════════════════════════════════════════════════════════════════════
    # SECTION 7: REMEDIATION ROADMAP
    # ═══════════════════════════════════════════════════════════════════════
    story.append(section_header("7. Remediation Roadmap"))
    story.append(Spacer(1, 0.4 * cm))

    if sorted_findings:
        rr_hdr_style = ParagraphStyle("rrh", fontName="Helvetica-Bold", fontSize=9, textColor=WHITE)
        rr_hdr = [Paragraph(h, rr_hdr_style) for h in
                  ["Priority", "ID", "Title", "Severity", "Recommended Action"]]
        rr_data = [rr_hdr]
        priority_labels = {
            "critical": "P1 — Immediate",
            "high":     "P2 — Urgent",
            "medium":   "P3 — Planned",
            "low":      "P4 — Scheduled",
            "informational": "P5 — Best Effort",
            "info":     "P5 — Best Effort",
        }
        for rank, f in enumerate(sorted_findings, 1):
            sev = (f.get("severity") or "").lower()
            pri = priority_labels.get(sev, "P5 — Best Effort")
            rec = _safe(f.get("recommendation"), "Review and apply security hardening.")
            rec_short = rec[:120] + "..." if len(rec) > 120 else rec
            rr_data.append([
                Paragraph(pri, styles["small_bold"]),
                Paragraph(_safe(f.get("finding_id")), styles["small_bold"]),
                Paragraph(_safe(f.get("title")), styles["small"]),
                sev_cell(sev),
                Paragraph(rec_short, styles["small"]),
            ])
        rr_tbl = Table(rr_data, colWidths=[3 * cm, 1.5 * cm, 4.5 * cm, 2.5 * cm, 5.5 * cm])
        rr_tbl.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), NAVY),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [WHITE, GREY_BG]),
            ("TOPPADDING", (0, 0), (-1, -1), 5),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ("LEFTPADDING", (0, 0), (-1, -1), 6),
            ("RIGHTPADDING", (0, 0), (-1, -1), 6),
            ("LINEBELOW", (0, 0), (-1, -1), 0.3, HexColor("#CFD8DC")),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ]))
        story.append(rr_tbl)
    else:
        story.append(Paragraph("No findings to remediate.", styles["normal"]))

    story.append(Spacer(1, 0.5 * cm))

    # ═══════════════════════════════════════════════════════════════════════
    # SECTION 8: CONCLUSION
    # ═══════════════════════════════════════════════════════════════════════
    story.append(section_header("8. Conclusion"))
    story.append(Spacer(1, 0.4 * cm))
    conclusion = _safe(report.get("conclusion"),
                       "This concludes the vulnerability assessment and penetration testing engagement.")
    for para in conclusion.split("\n"):
        if para.strip():
            story.append(Paragraph(para.strip(), styles["normal"]))
            story.append(Spacer(1, 0.15 * cm))

    story.append(Spacer(1, 0.5 * cm))

    # ═══════════════════════════════════════════════════════════════════════
    # SECTION 9: APPENDICES
    # ═══════════════════════════════════════════════════════════════════════
    appendices = report.get("appendices")
    if appendices and str(appendices).strip():
        story.append(section_header("9. Appendices"))
        story.append(Spacer(1, 0.4 * cm))
        for para in str(appendices).split("\n"):
            if para.strip():
                story.append(Paragraph(para.strip(), styles["normal"]))
                story.append(Spacer(1, 0.15 * cm))

    doc.build(story)
    return buf.getvalue()


# ══════════════════════════════════════════════════════════════════════════════
# DOCX GENERATION
# ══════════════════════════════════════════════════════════════════════════════

def _set_cell_shading(cell, hex_color: str):
    """Apply background colour to a docx table cell via CT_Shd XML."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tcPr.append(shd)


def _cell_para_style(cell, bold=False, font_size=10, color_hex="000000", center=False):
    for para in cell.paragraphs:
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        if center:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.bold = bold
            run.font.size = Pt(font_size)
            run.font.color.rgb = RGBColor.from_string(color_hex)


SEV_DOCX_COLORS = {
    "critical":      "C62828",
    "high":          "E64A19",
    "medium":        "F9A825",
    "low":           "2E7D32",
    "informational": "1565C0",
    "info":          "1565C0",
}

RETEST_DOCX_COLORS = {
    "passed":         "2E7D32",
    "failed":         "C62828",
    "pending":        "757575",
    "not_applicable": "455A64",
}


def generate_docx(report: Dict, findings: List[Dict], client_name: str) -> bytes:
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    sorted_findings = sorted(findings, key=_sev_key)
    sev_counts = _sev_counts(findings)
    scope = {}
    methodology = {}
    try:
        scope = json.loads(report.get("scope_json") or "{}")
    except Exception:
        pass
    try:
        methodology = json.loads(report.get("methodology_json") or "{}")
    except Exception:
        pass

    report_title = _safe(report.get("title"), "VAPT Report")
    version_str = _safe(report.get("version"), "1.0")
    classification = _safe(report.get("classification"), "Confidential")

    # ── Style helpers ─────────────────────────────────────────────────────
    def _add_heading(text: str, level: int = 1, color_hex: str = "1A237E"):
        p = doc.add_heading(text, level=level)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        for run in p.runs:
            run.font.color.rgb = RGBColor.from_string(color_hex)
            run.font.bold = True
        return p

    def _add_para(text: str, bold: bool = False, size: int = 10, color: str = "37474F"):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor.from_string(color)
        run.font.name = "Calibri"
        p.paragraph_format.space_after = Pt(4)
        return p

    def _add_section_header(text: str):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.style = "Table Grid"
        cell = tbl.cell(0, 0)
        _set_cell_shading(cell, "1A237E")
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.name = "Calibri"
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        doc.add_paragraph()

    # ── Cover page ────────────────────────────────────────────────────────
    # Title block
    cover_tbl = doc.add_table(rows=1, cols=1)
    cover_tbl.style = "Table Grid"
    cell = cover_tbl.cell(0, 0)
    _set_cell_shading(cell, "0D1B4B")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("MONITARA AI\n")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x42, 0xA5, 0xF5)
    run.font.name = "Calibri"
    run2 = p.add_run("Cybersecurity Platform\n\n")
    run2.font.size = Pt(11)
    run2.font.color.rgb = RGBColor(0x90, 0xCA, 0xF9)
    run2.font.name = "Calibri"
    run3 = p.add_run(report_title.upper() + "\n")
    run3.bold = True
    run3.font.size = Pt(22)
    run3.font.color.rgb = RGBColor(255, 255, 255)
    run3.font.name = "Calibri"
    run4 = p.add_run("Vulnerability Assessment & Penetration Testing Report\n")
    run4.font.size = Pt(12)
    run4.font.color.rgb = RGBColor(0xB0, 0xBE, 0xC5)
    run4.font.name = "Calibri"

    doc.add_paragraph()

    # Metadata table
    meta_tbl = doc.add_table(rows=7, cols=2)
    meta_tbl.style = "Table Grid"
    meta_rows = [
        ("Client", _safe(client_name)),
        ("Classification", classification),
        ("Version", f"v{version_str}"),
        ("Report Date", _fmt_date(report.get("report_date"))),
        ("Prepared By", _safe(report.get("prepared_by"))),
        ("Reviewed By", _safe(report.get("reviewed_by"))),
        ("Status", _safe(report.get("status"), "Draft").capitalize()),
    ]
    for i, (label, value) in enumerate(meta_rows):
        lbl_cell = meta_tbl.cell(i, 0)
        val_cell = meta_tbl.cell(i, 1)
        _set_cell_shading(lbl_cell, "ECEFF1")
        lp = lbl_cell.paragraphs[0]
        lr = lp.add_run(label)
        lr.bold = True
        lr.font.size = Pt(10)
        lr.font.color.rgb = RGBColor(0x54, 0x6E, 0x7A)
        vp = val_cell.paragraphs[0]
        vr = vp.add_run(value)
        vr.font.size = Pt(10)

    doc.add_paragraph()
    _add_para("CONFIDENTIAL — This document contains sensitive security information. "
              "Distribution is restricted to authorised personnel only.",
              size=9, color="90A4AE")
    doc.add_page_break()

    # ── Section 1: Executive Summary ──────────────────────────────────────
    _add_section_header("1. Executive Summary")

    # Severity stats table
    sev_display = [
        ("CRITICAL", "critical"),
        ("HIGH", "high"),
        ("MEDIUM", "medium"),
        ("LOW", "low"),
        ("INFO", "informational"),
        ("TOTAL", None),
    ]
    sev_tbl = doc.add_table(rows=2, cols=6)
    sev_tbl.style = "Table Grid"
    for ci, (label, key) in enumerate(sev_display):
        top_cell = sev_tbl.cell(0, ci)
        bot_cell = sev_tbl.cell(1, ci)
        bg = SEV_DOCX_COLORS.get(key, "1A237E") if key else "1A237E"
        _set_cell_shading(top_cell, bg)
        _set_cell_shading(bot_cell, bg)
        tp = top_cell.paragraphs[0]
        tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tr = tp.add_run(label)
        tr.bold = True
        tr.font.size = Pt(8)
        tr.font.color.rgb = RGBColor(255, 255, 255)
        bp = bot_cell.paragraphs[0]
        bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        count = sev_counts.get(key, 0) if key else sum(sev_counts.values())
        br = bp.add_run(str(count))
        br.bold = True
        br.font.size = Pt(18)
        br.font.color.rgb = RGBColor(255, 255, 255)

    doc.add_paragraph()
    exec_summary = _safe(report.get("executive_summary"), "No executive summary provided.")
    _add_para(exec_summary)
    doc.add_paragraph()

    # ── Section 2: Scope ──────────────────────────────────────────────────
    _add_section_header("2. Scope of Engagement")

    scope_tbl = doc.add_table(rows=3, cols=2)
    scope_tbl.style = "Table Grid"
    scope_meta = [
        ("Engagement Type", _safe(scope.get("engagement_type"))),
        ("Period Start", _safe(scope.get("period_start"))),
        ("Period End", _safe(scope.get("period_end"))),
    ]
    for i, (label, value) in enumerate(scope_meta):
        lc = scope_tbl.cell(i, 0)
        vc = scope_tbl.cell(i, 1)
        _set_cell_shading(lc, "ECEFF1")
        lp = lc.paragraphs[0]
        lr = lp.add_run(label)
        lr.bold = True
        lr.font.size = Pt(10)
        vp = vc.paragraphs[0]
        vr = vp.add_run(value)
        vr.font.size = Pt(10)

    doc.add_paragraph()
    in_scope = scope.get("in_scope") or []
    if in_scope:
        _add_para("In-Scope Assets:", bold=True, color="1565C0")
        for item in in_scope:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(_safe(item))

    out_scope = scope.get("out_of_scope") or []
    if out_scope:
        _add_para("Out-of-Scope Items:", bold=True, color="1565C0")
        for item in out_scope:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(_safe(item))

    doc.add_paragraph()

    # ── Section 3: Methodology ─────────────────────────────────────────────
    _add_section_header("3. Methodology")

    phases = methodology.get("phases") or []
    if phases:
        _add_para("Testing Phases:", bold=True, color="1565C0")
        ph_tbl = doc.add_table(rows=len(phases) + 1, cols=3)
        ph_tbl.style = "Table Grid"
        for ci, hdr in enumerate(["#", "Phase", "Description"]):
            c = ph_tbl.cell(0, ci)
            _set_cell_shading(c, "1A237E")
            p = c.paragraphs[0]
            r = p.add_run(hdr)
            r.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)
        for i, phase in enumerate(phases):
            ph_tbl.cell(i + 1, 0).paragraphs[0].add_run(str(i + 1))
            ph_tbl.cell(i + 1, 1).paragraphs[0].add_run(_safe(phase.get("name")))
            ph_tbl.cell(i + 1, 2).paragraphs[0].add_run(_safe(phase.get("description")))
        doc.add_paragraph()

    tools = methodology.get("tools") or []
    if tools:
        _add_para("Tools Used:", bold=True, color="1565C0")
        _add_para(" | ".join(tools))

    standards = methodology.get("standards") or []
    if standards:
        _add_para("Standards Referenced:", bold=True, color="1565C0")
        for std in standards:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(_safe(std))

    doc.add_paragraph()

    # ── Section 4: Severity Matrix ─────────────────────────────────────────
    _add_section_header("4. Severity Rating Matrix")

    matrix_data = [
        ("CRITICAL", "9.0–10.0", "C62828", "Remote unauthenticated RCE, complete system compromise",
         "Immediate — 24h"),
        ("HIGH", "7.0–8.9", "E64A19", "Significant risk requiring urgent remediation",
         "Urgent — 72h"),
        ("MEDIUM", "4.0–6.9", "F9A825", "Moderate risk requiring planned remediation",
         "Planned — 30 days"),
        ("LOW", "0.1–3.9", "2E7D32", "Minor risk with limited exposure",
         "Scheduled — 90 days"),
        ("INFORMATIONAL", "N/A", "1565C0", "Best practice observations, no immediate risk",
         "Best effort"),
    ]
    m_tbl = doc.add_table(rows=len(matrix_data) + 1, cols=4)
    m_tbl.style = "Table Grid"
    for ci, hdr in enumerate(["Rating", "CVSS", "Description", "SLA"]):
        c = m_tbl.cell(0, ci)
        _set_cell_shading(c, "1A237E")
        p = c.paragraphs[0]
        r = p.add_run(hdr)
        r.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
    for i, (label, cvss, color, desc, sla) in enumerate(matrix_data):
        cells = [m_tbl.cell(i + 1, ci) for ci in range(4)]
        _set_cell_shading(cells[0], color)
        p = cells[0].paragraphs[0]
        r = p.add_run(label)
        r.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        cells[1].paragraphs[0].add_run(cvss)
        cells[2].paragraphs[0].add_run(desc)
        cells[3].paragraphs[0].add_run(sla)
    doc.add_paragraph()

    # ── Section 5: Findings Summary ────────────────────────────────────────
    _add_section_header("5. Findings Summary")

    if sorted_findings:
        fs_tbl = doc.add_table(rows=len(sorted_findings) + 1, cols=5)
        fs_tbl.style = "Table Grid"
        for ci, hdr in enumerate(["ID", "Title", "Severity", "Asset", "Retest"]):
            c = fs_tbl.cell(0, ci)
            _set_cell_shading(c, "1A237E")
            p = c.paragraphs[0]
            r = p.add_run(hdr)
            r.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)
        for i, f in enumerate(sorted_findings):
            sev = (f.get("severity") or "").lower()
            rs = (f.get("retest_status") or "pending").lower()
            row_cells = [fs_tbl.cell(i + 1, ci) for ci in range(5)]
            row_cells[0].paragraphs[0].add_run(_safe(f.get("finding_id")))
            row_cells[1].paragraphs[0].add_run(_safe(f.get("title")))
            sev_c = row_cells[2]
            _set_cell_shading(sev_c, SEV_DOCX_COLORS.get(sev, "757575"))
            sp = sev_c.paragraphs[0]
            sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sr = sp.add_run(sev.capitalize())
            sr.font.color.rgb = RGBColor(255, 255, 255)
            sr.bold = True
            row_cells[3].paragraphs[0].add_run(_safe(f.get("affected_asset")))
            rs_c = row_cells[4]
            _set_cell_shading(rs_c, RETEST_DOCX_COLORS.get(rs, "757575"))
            rp = rs_c.paragraphs[0]
            rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            rr2 = rp.add_run(rs.replace("_", " ").title())
            rr2.font.color.rgb = RGBColor(255, 255, 255)
        doc.add_paragraph()
    else:
        _add_para("No findings recorded.")

    # ── Section 6: Detailed Findings ───────────────────────────────────────
    _add_section_header("6. Detailed Findings")

    for fi, f in enumerate(sorted_findings):
        sev = (f.get("severity") or "").lower()
        fid = _safe(f.get("finding_id"), f"F-{fi + 1:02d}")
        ftitle = _safe(f.get("title"))
        rs = (f.get("retest_status") or "pending").lower()

        # Color-coded header
        f_hdr_tbl = doc.add_table(rows=1, cols=1)
        f_hdr_tbl.style = "Table Grid"
        hdr_cell = f_hdr_tbl.cell(0, 0)
        _set_cell_shading(hdr_cell, SEV_DOCX_COLORS.get(sev, "757575"))
        hp = hdr_cell.paragraphs[0]
        hr2 = hp.add_run(f"{fid}  —  {ftitle}")
        hr2.bold = True
        hr2.font.size = Pt(12)
        hr2.font.color.rgb = RGBColor(255, 255, 255)

        # Meta row
        meta_tbl = doc.add_table(rows=1, cols=6)
        meta_tbl.style = "Table Grid"
        meta_pairs = [
            ("Severity", sev.capitalize()),
            ("Affected Asset", _safe(f.get("affected_asset"))),
            ("Retest Status", rs.replace("_", " ").title()),
        ]
        _set_cell_shading(meta_tbl.cell(0, 0), "ECEFF1")
        _set_cell_shading(meta_tbl.cell(0, 2), "ECEFF1")
        _set_cell_shading(meta_tbl.cell(0, 4), "ECEFF1")
        for pair_idx, (lbl, val) in enumerate(meta_pairs):
            lc = meta_tbl.cell(0, pair_idx * 2)
            vc = meta_tbl.cell(0, pair_idx * 2 + 1)
            lr = lc.paragraphs[0].add_run(lbl + ":")
            lr.bold = True
            lr.font.size = Pt(9)
            vr = vc.paragraphs[0].add_run(val)
            vr.font.size = Pt(9)

        def _field(label: str, value: str):
            if not value or value == "—":
                return
            lp2 = doc.add_paragraph()
            lr3 = lp2.add_run(label)
            lr3.bold = True
            lr3.font.color.rgb = RGBColor(0x15, 0x65, 0xC0)
            _add_para(value)

        _field("Description", _safe(f.get("description")))
        _field("Business Impact", _safe(f.get("impact")))
        _field("Evidence", _safe(f.get("evidence")))
        _field("Reproduction Steps", _safe(f.get("reproduction_steps")))
        _field("Recommendation", _safe(f.get("recommendation")))
        _field("References", _safe(f.get("references")))

        rn = f.get("retest_notes")
        if rn and rn.strip():
            rn_tbl = doc.add_table(rows=2, cols=1)
            rn_tbl.style = "Table Grid"
            rn_hdr_c = rn_tbl.cell(0, 0)
            rn_body_c = rn_tbl.cell(1, 0)
            rn_color = RETEST_DOCX_COLORS.get(rs, "757575")
            _set_cell_shading(rn_hdr_c, rn_color)
            _set_cell_shading(rn_body_c, rn_color)
            rh = rn_hdr_c.paragraphs[0].add_run(f"Retest Notes ({rs.replace('_', ' ').title()})")
            rh.bold = True
            rh.font.color.rgb = RGBColor(255, 255, 255)
            rb = rn_body_c.paragraphs[0].add_run(rn)
            rb.font.color.rgb = RGBColor(255, 255, 255)
            rb.font.size = Pt(9)

        doc.add_paragraph()

    # ── Section 7: Remediation Roadmap ─────────────────────────────────────
    _add_section_header("7. Remediation Roadmap")

    if sorted_findings:
        rr_tbl = doc.add_table(rows=len(sorted_findings) + 1, cols=4)
        rr_tbl.style = "Table Grid"
        for ci, hdr in enumerate(["Priority", "ID", "Title/Severity", "Recommended Action"]):
            c = rr_tbl.cell(0, ci)
            _set_cell_shading(c, "1A237E")
            p = c.paragraphs[0]
            r = p.add_run(hdr)
            r.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)
        priority_labels = {
            "critical": "P1 — Immediate",
            "high": "P2 — Urgent",
            "medium": "P3 — Planned",
            "low": "P4 — Scheduled",
            "informational": "P5 — Best Effort",
            "info": "P5 — Best Effort",
        }
        for rank, f in enumerate(sorted_findings, 1):
            sev = (f.get("severity") or "").lower()
            pri = priority_labels.get(sev, "P5")
            rec = _safe(f.get("recommendation"), "Review and apply security hardening.")
            rec_short = rec[:200] + "..." if len(rec) > 200 else rec
            row_cells = [rr_tbl.cell(rank, ci) for ci in range(4)]
            row_cells[0].paragraphs[0].add_run(pri)
            row_cells[1].paragraphs[0].add_run(_safe(f.get("finding_id")))
            title_c = row_cells[2]
            tp = title_c.paragraphs[0]
            tp.add_run(_safe(f.get("title")) + " ")
            _set_cell_shading(title_c, SEV_DOCX_COLORS.get(sev, "757575"))
            tr2 = tp.runs[-1]
            tr2.font.color.rgb = RGBColor(255, 255, 255)
            row_cells[3].paragraphs[0].add_run(rec_short)
        doc.add_paragraph()
    else:
        _add_para("No findings to remediate.")

    # ── Section 8: Conclusion ──────────────────────────────────────────────
    _add_section_header("8. Conclusion")
    conclusion = _safe(report.get("conclusion"),
                       "This concludes the vulnerability assessment and penetration testing engagement.")
    _add_para(conclusion)
    doc.add_paragraph()

    # ── Section 9: Appendices ──────────────────────────────────────────────
    appendices = report.get("appendices")
    if appendices and str(appendices).strip():
        _add_section_header("9. Appendices")
        _add_para(str(appendices))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ══════════════════════════════════════════════════════════════════════════════
# REMEDIATION PLAN PDF
# ══════════════════════════════════════════════════════════════════════════════

def generate_remediation_pdf(report: Dict, findings: List[Dict], client_name: str) -> bytes:
    buf = io.BytesIO()
    styles = _build_pdf_styles()
    sorted_findings = sorted(findings, key=_sev_key)
    sev_counts = _sev_counts(findings)

    PAGE_W, PAGE_H = A4
    MARGIN = 2 * cm
    report_title = _safe(report.get("title"), "VAPT Report")
    version_str = _safe(report.get("version"), "1.0")
    classification = _safe(report.get("classification"), "Confidential")

    def _cover_template(canvas, doc):
        canvas.saveState()
        canvas.setFillColor(HexColor("#1B2A4A"))
        canvas.rect(0, 0, PAGE_W, PAGE_H, fill=1, stroke=0)
        canvas.setFillColor(HexColor("#C62828"))
        canvas.rect(0, PAGE_H - 0.8 * cm, PAGE_W, 0.8 * cm, fill=1, stroke=0)
        canvas.setFillColor(HexColor("#C62828"))
        canvas.rect(0, 0, PAGE_W, 0.8 * cm, fill=1, stroke=0)
        canvas.restoreState()

    def _content_template(canvas, doc):
        canvas.saveState()
        canvas.setFillColor(HexColor("#C62828"))
        canvas.rect(MARGIN, PAGE_H - MARGIN + 0.2 * cm, PAGE_W - 2 * MARGIN, 0.05 * cm, fill=1, stroke=0)
        canvas.setFont("Helvetica", 8)
        canvas.setFillColor(DARK_TEXT)
        canvas.drawString(MARGIN, PAGE_H - MARGIN + 0.45 * cm, "REMEDIATION ACTION PLAN")
        canvas.drawRightString(PAGE_W - MARGIN, PAGE_H - MARGIN + 0.45 * cm,
                               f"{classification}  |  Page {doc.page}")
        canvas.setFillColor(HexColor("#C62828"))
        canvas.rect(MARGIN, MARGIN - 0.4 * cm, PAGE_W - 2 * MARGIN, 0.05 * cm, fill=1, stroke=0)
        canvas.setFont("Helvetica", 8)
        canvas.setFillColor(DARK_TEXT)
        canvas.drawString(MARGIN, MARGIN - 0.6 * cm, client_name)
        canvas.drawRightString(PAGE_W - MARGIN, MARGIN - 0.6 * cm, f"Version {version_str}")
        canvas.restoreState()

    cover_frame = Frame(0, 0, PAGE_W, PAGE_H, leftPadding=2.5 * cm, rightPadding=1.5 * cm,
                        topPadding=2 * cm, bottomPadding=2 * cm)
    content_frame = Frame(MARGIN, MARGIN, PAGE_W - 2 * MARGIN, PAGE_H - 2 * MARGIN,
                          topPadding=0.8 * cm, bottomPadding=0.8 * cm)

    doc = BaseDocTemplate(buf, pagesize=A4, leftMargin=MARGIN, rightMargin=MARGIN,
                          topMargin=MARGIN, bottomMargin=MARGIN)
    doc.addPageTemplates([
        PageTemplate(id="Cover", frames=[cover_frame], onPage=_cover_template),
        PageTemplate(id="Content", frames=[content_frame], onPage=_content_template),
    ])

    story = []

    def section_header(text):
        tbl = Table([[Paragraph(text, styles["section"])]], colWidths=[PAGE_W - 2 * MARGIN])
        tbl.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, -1), HexColor("#C62828")),
            ("LEFTPADDING", (0, 0), (-1, -1), 10),
            ("RIGHTPADDING", (0, 0), (-1, -1), 10),
            ("TOPPADDING", (0, 0), (-1, -1), 6),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
        ]))
        return tbl

    # Cover
    story.append(Spacer(1, 3 * cm))
    story.append(Paragraph("REMEDIATION ACTION PLAN", ParagraphStyle(
        "rap_title", fontName="Helvetica-Bold", fontSize=26, textColor=WHITE,
        alignment=TA_CENTER, leading=32)))
    story.append(Spacer(1, 0.5 * cm))
    story.append(Paragraph("Based on VAPT Engagement Findings", ParagraphStyle(
        "rap_sub", fontName="Helvetica", fontSize=13, textColor=HexColor("#90A4AE"),
        alignment=TA_CENTER)))
    story.append(Spacer(1, 2 * cm))
    story.append(HRFlowable(width="70%", thickness=2, color=HexColor("#C62828"), hAlign="CENTER"))
    story.append(Spacer(1, 1.5 * cm))

    cover_style_bold = ParagraphStyle("cb", fontName="Helvetica-Bold", fontSize=10,
                                      textColor=HexColor("#90CAF9"))
    cover_style_normal = ParagraphStyle("cn", fontName="Helvetica", fontSize=10, textColor=WHITE)
    cover_data = [
        ["Client", _safe(client_name)],
        ["Original Report", report_title],
        ["Report Version", f"v{version_str}"],
        ["Date", _fmt_date(report.get("report_date"))],
        ["Classification", classification],
    ]
    cover_tbl_data = [[Paragraph(r[0], cover_style_bold), Paragraph(r[1], cover_style_normal)]
                      for r in cover_data]
    ctbl = Table(cover_tbl_data, colWidths=[4 * cm, 10 * cm])
    ctbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), HexColor("#0A1340")),
        ("LINEBELOW", (0, 0), (-1, -2), 0.5, HexColor("#1E3A8A")),
        ("TOPPADDING", (0, 0), (-1, -1), 7),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
        ("LEFTPADDING", (0, 0), (-1, -1), 10),
        ("RIGHTPADDING", (0, 0), (-1, -1), 10),
    ]))
    story.append(ctbl)
    story.append(PageBreak())

    from reportlab.platypus import NextPageTemplate
    story.append(NextPageTemplate("Content"))

    # Section 1: Executive Overview
    story.append(section_header("1. Executive Overview"))
    story.append(Spacer(1, 0.4 * cm))

    sev_label_s = ParagraphStyle("sl2", fontName="Helvetica", fontSize=9, textColor=WHITE, alignment=TA_CENTER)
    sev_count_s = ParagraphStyle("sc2", fontName="Helvetica-Bold", fontSize=20, textColor=WHITE, alignment=TA_CENTER)
    sev_disp = [("CRITICAL", "critical"), ("HIGH", "high"), ("MEDIUM", "medium"),
                ("LOW", "low"), ("INFO", "informational"), ("TOTAL", None)]
    sev_row = []
    for label, key in sev_disp:
        bg = SEV_COLORS.get(key, NAVY) if key else NAVY
        count = sev_counts.get(key, 0) if key else sum(sev_counts.values())
        cell = Table([[Paragraph(label, sev_label_s)], [Paragraph(str(count), sev_count_s)]],
                     colWidths=[2.5 * cm])
        cell.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, -1), bg),
            ("TOPPADDING", (0, 0), (-1, -1), 6),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ("LEFTPADDING", (0, 0), (-1, -1), 4),
            ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ]))
        sev_row.append(cell)
    story.append(Table([sev_row], colWidths=[2.5 * cm] * 6))
    story.append(Spacer(1, 0.5 * cm))
    exec_text = _safe(report.get("executive_summary"),
                      "This document provides a prioritised remediation plan based on the findings of the VAPT engagement.")
    story.append(Paragraph(exec_text[:500], styles["normal"]))
    story.append(Spacer(1, 0.5 * cm))

    # Section 2: Prioritized Action Table
    story.append(section_header("2. Prioritised Action Table"))
    story.append(Spacer(1, 0.4 * cm))

    effort_map = {"critical": "High", "high": "High", "medium": "Medium", "low": "Low", "informational": "Low"}
    owner_map = {"critical": "Security Team + CTO", "high": "Security Team", "medium": "Dev/Ops Team",
                 "low": "Dev Team", "informational": "Dev Team"}
    priority_labels = {"critical": "P1 — Immediate", "high": "P2 — Urgent",
                       "medium": "P3 — Planned", "low": "P4 — Scheduled",
                       "informational": "P5 — Best Effort", "info": "P5 — Best Effort"}

    def sev_cell(sev: str) -> Table:
        s = sev.lower()
        bg = SEV_COLORS.get(s, HexColor("#757575"))
        tc = SEV_TEXT_COLORS.get(s, WHITE)
        label = sev.capitalize()
        p = ParagraphStyle("sev_cell2", fontName="Helvetica-Bold", fontSize=9,
                           textColor=tc, alignment=TA_CENTER)
        t = Table([[Paragraph(label, p)]], colWidths=[2 * cm])
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, -1), bg),
            ("TOPPADDING", (0, 0), (-1, -1), 3),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ("LEFTPADDING", (0, 0), (-1, -1), 2),
            ("RIGHTPADDING", (0, 0), (-1, -1), 2),
        ]))
        return t

    pat_hdr = [Paragraph(h, ParagraphStyle("ph2", fontName="Helvetica-Bold", fontSize=9, textColor=WHITE))
               for h in ["Rank", "ID", "Title", "Severity", "Effort", "Owner", "Status"]]
    pat_data = [pat_hdr]
    for rank, f in enumerate(sorted_findings, 1):
        sev = (f.get("severity") or "").lower()
        rs = (f.get("retest_status") or "pending").lower()
        pat_data.append([
            Paragraph(str(rank), styles["small_bold"]),
            Paragraph(_safe(f.get("finding_id")), styles["small_bold"]),
            Paragraph(_safe(f.get("title")), styles["small"]),
            sev_cell(sev) if sev else Paragraph("—", styles["small"]),
            Paragraph(effort_map.get(sev, "Medium"), styles["small"]),
            Paragraph(owner_map.get(sev, "Dev Team"), styles["small"]),
            Paragraph(rs.replace("_", " ").title(), styles["small"]),
        ])

    pat_tbl = Table(pat_data, colWidths=[1 * cm, 1.3 * cm, 4.5 * cm, 2 * cm, 1.5 * cm, 3 * cm, 2.5 * cm])
    pat_tbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), HexColor("#C62828")),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [WHITE, GREY_BG]),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ("LEFTPADDING", (0, 0), (-1, -1), 5),
        ("RIGHTPADDING", (0, 0), (-1, -1), 5),
        ("LINEBELOW", (0, 0), (-1, -1), 0.3, HexColor("#CFD8DC")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
    ]))
    story.append(pat_tbl)
    story.append(Spacer(1, 0.5 * cm))

    # Section 3: Detailed Remediation Steps
    story.append(section_header("3. Detailed Remediation Steps"))
    story.append(Spacer(1, 0.4 * cm))

    for fi, f in enumerate(sorted_findings):
        sev = (f.get("severity") or "").lower()
        fid = _safe(f.get("finding_id"), f"F-{fi + 1:02d}")
        bg = SEV_COLORS.get(sev, HexColor("#757575"))
        tc2 = SEV_TEXT_COLORS.get(sev, WHITE)
        rs = (f.get("retest_status") or "pending").lower()

        elems = []

        hdr_style = ParagraphStyle("fdh2", fontName="Helvetica-Bold", fontSize=11, textColor=tc2)
        fh_tbl = Table([[Paragraph(f"{fid}  —  {_safe(f.get('title'))}", hdr_style)]],
                        colWidths=[PAGE_W - 2 * MARGIN])
        fh_tbl.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, -1), bg),
            ("TOPPADDING", (0, 0), (-1, -1), 8),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
            ("LEFTPADDING", (0, 0), (-1, -1), 10),
            ("RIGHTPADDING", (0, 0), (-1, -1), 10),
        ]))
        elems.append(fh_tbl)
        elems.append(Spacer(1, 0.2 * cm))

        # Context
        if f.get("description") or f.get("impact"):
            elems.append(Paragraph("Context", styles["subsection"]))
            if f.get("description"):
                elems.append(Paragraph(f"What was found: {_safe(f.get('description'))[:400]}", styles["normal"]))
            if f.get("impact"):
                elems.append(Paragraph(f"Why it matters: {_safe(f.get('impact'))[:300]}", styles["normal"]))
            elems.append(Spacer(1, 0.2 * cm))

        # Immediate Actions
        rec = _safe(f.get("recommendation"))
        if rec and rec != "—":
            elems.append(Paragraph("Immediate Actions", styles["subsection"]))
            for step_i, line in enumerate(rec.split("\n"), 1):
                if line.strip():
                    elems.append(Paragraph(f"{step_i}. {line.strip()}", styles["bullet"]))
            elems.append(Spacer(1, 0.2 * cm))

        # Verification steps
        elems.append(Paragraph("Verification Steps", styles["subsection"]))
        elems.append(Paragraph("1. Apply the recommended fix in a test environment.", styles["bullet"]))
        elems.append(Paragraph("2. Re-run the relevant test case or scan to confirm the vulnerability is remediated.", styles["bullet"]))
        elems.append(Paragraph("3. Document evidence of the fix (screenshots, config exports, test results).", styles["bullet"]))
        elems.append(Paragraph("4. Submit for retest validation by the security team.", styles["bullet"]))
        elems.append(Spacer(1, 0.2 * cm))

        # Effort
        effort_label = effort_map.get(sev, "Medium")
        elems.append(Paragraph(f"Estimated Effort: {effort_label}  |  Priority: {priority_labels.get(sev, 'P3')}",
                               styles["label"]))

        if fi < len(sorted_findings) - 1:
            elems.append(HRFlowable(width="100%", thickness=1, color=HexColor("#CFD8DC")))
            elems.append(Spacer(1, 0.3 * cm))

        try:
            story.append(KeepTogether(elems[:6]))
            for el in elems[6:]:
                story.append(el)
        except Exception:
            for el in elems:
                story.append(el)

    story.append(Spacer(1, 0.5 * cm))

    # Section 4: Verification Checklist
    story.append(section_header("4. Verification Checklist"))
    story.append(Spacer(1, 0.3 * cm))

    chk_hdr = [Paragraph(h, ParagraphStyle("chkh", fontName="Helvetica-Bold", fontSize=9, textColor=WHITE))
               for h in ["ID", "Title", "Severity", "Fix Applied", "Verified", "Closed"]]
    chk_data = [chk_hdr]
    for f in sorted_findings:
        sev = (f.get("severity") or "").lower()
        chk_data.append([
            Paragraph(_safe(f.get("finding_id")), styles["small_bold"]),
            Paragraph(_safe(f.get("title")), styles["small"]),
            sev_cell(sev),
            Paragraph("☐", styles["center"]),
            Paragraph("☐", styles["center"]),
            Paragraph("☐", styles["center"]),
        ])
    chk_tbl = Table(chk_data, colWidths=[1.5 * cm, 6.5 * cm, 2.5 * cm, 1.8 * cm, 1.8 * cm, 1.8 * cm])
    chk_tbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), HexColor("#C62828")),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [WHITE, GREY_BG]),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("LINEBELOW", (0, 0), (-1, -1), 0.3, HexColor("#CFD8DC")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("ALIGN", (3, 0), (-1, -1), "CENTER"),
    ]))
    story.append(chk_tbl)
    story.append(Spacer(1, 0.5 * cm))

    # Section 5: Sign-off
    story.append(section_header("5. Sign-off"))
    story.append(Spacer(1, 0.3 * cm))
    signoff_data = [
        [Paragraph(h, ParagraphStyle("soh", fontName="Helvetica-Bold", fontSize=9, textColor=WHITE))
         for h in ["Role", "Name", "Signature", "Date"]],
        ["Prepared By", _safe(report.get("prepared_by")), "", ""],
        ["Reviewed By", _safe(report.get("reviewed_by")), "", ""],
        ["Client Sign-off", "", "", ""],
    ]
    for i in range(1, len(signoff_data)):
        for j in range(len(signoff_data[i])):
            if not isinstance(signoff_data[i][j], Paragraph):
                signoff_data[i][j] = Paragraph(str(signoff_data[i][j]), styles["normal"])
    signoff_tbl = Table(signoff_data, colWidths=[4 * cm, 5 * cm, 4 * cm, 4 * cm])
    signoff_tbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), HexColor("#C62828")),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [WHITE, GREY_BG]),
        ("LINEBELOW", (0, 0), (-1, -1), 0.5, HexColor("#CFD8DC")),
        ("TOPPADDING", (0, 0), (-1, -1), 10),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 10),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
    ]))
    story.append(signoff_tbl)

    doc.build(story)
    return buf.getvalue()


# ══════════════════════════════════════════════════════════════════════════════
# REMEDIATION PLAN DOCX
# ══════════════════════════════════════════════════════════════════════════════

def generate_remediation_docx(report: Dict, findings: List[Dict], client_name: str) -> bytes:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    sorted_findings = sorted(findings, key=_sev_key)
    sev_counts = _sev_counts(findings)
    version_str = _safe(report.get("version"), "1.0")
    report_title = _safe(report.get("title"), "VAPT Report")
    classification = _safe(report.get("classification"), "Confidential")

    priority_labels = {"critical": "P1 — Immediate", "high": "P2 — Urgent",
                       "medium": "P3 — Planned", "low": "P4 — Scheduled",
                       "informational": "P5 — Best Effort", "info": "P5 — Best Effort"}
    effort_map = {"critical": "High", "high": "High", "medium": "Medium", "low": "Low", "informational": "Low"}
    owner_map = {"critical": "Security Team + CTO", "high": "Security Team",
                 "medium": "Dev/Ops Team", "low": "Dev Team", "informational": "Dev Team"}

    def _add_section_header(text: str, color: str = "C62828"):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.style = "Table Grid"
        cell = tbl.cell(0, 0)
        _set_cell_shading(cell, color)
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(255, 255, 255)
        doc.add_paragraph()

    def _add_para(text: str, bold: bool = False, size: int = 10, color: str = "37474F"):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor.from_string(color)
        run.font.name = "Calibri"
        p.paragraph_format.space_after = Pt(4)
        return p

    # Cover
    cover_tbl = doc.add_table(rows=1, cols=1)
    cover_tbl.style = "Table Grid"
    cell = cover_tbl.cell(0, 0)
    _set_cell_shading(cell, "1B2A4A")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("REMEDIATION ACTION PLAN\n")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor(255, 255, 255)
    r2 = p.add_run("Based on VAPT Engagement Findings\n")
    r2.font.size = Pt(11)
    r2.font.color.rgb = RGBColor(0xB0, 0xBE, 0xC5)
    doc.add_paragraph()

    meta_tbl = doc.add_table(rows=5, cols=2)
    meta_tbl.style = "Table Grid"
    for i, (label, value) in enumerate([
        ("Client", _safe(client_name)),
        ("Original Report", report_title),
        ("Version", f"v{version_str}"),
        ("Date", _fmt_date(report.get("report_date"))),
        ("Classification", classification),
    ]):
        lc = meta_tbl.cell(i, 0)
        vc = meta_tbl.cell(i, 1)
        _set_cell_shading(lc, "ECEFF1")
        lr = lc.paragraphs[0].add_run(label)
        lr.bold = True
        lr.font.size = Pt(10)
        vc.paragraphs[0].add_run(value)
    doc.add_page_break()

    # Section 1: Executive Overview
    _add_section_header("1. Executive Overview")
    sev_tbl = doc.add_table(rows=2, cols=6)
    sev_tbl.style = "Table Grid"
    sev_disp = [("CRITICAL", "critical"), ("HIGH", "high"), ("MEDIUM", "medium"),
                ("LOW", "low"), ("INFO", "informational"), ("TOTAL", None)]
    for ci, (label, key) in enumerate(sev_disp):
        bg = SEV_DOCX_COLORS.get(key, "1A237E") if key else "1A237E"
        top_c = sev_tbl.cell(0, ci)
        bot_c = sev_tbl.cell(1, ci)
        _set_cell_shading(top_c, bg)
        _set_cell_shading(bot_c, bg)
        tp = top_c.paragraphs[0]
        tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tr = tp.add_run(label)
        tr.bold = True
        tr.font.size = Pt(8)
        tr.font.color.rgb = RGBColor(255, 255, 255)
        bp = bot_c.paragraphs[0]
        bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        count = sev_counts.get(key, 0) if key else sum(sev_counts.values())
        br = bp.add_run(str(count))
        br.bold = True
        br.font.size = Pt(16)
        br.font.color.rgb = RGBColor(255, 255, 255)
    doc.add_paragraph()

    # Section 2: Prioritized Action Table
    _add_section_header("2. Prioritised Action Table")
    if sorted_findings:
        rr_tbl = doc.add_table(rows=len(sorted_findings) + 1, cols=5)
        rr_tbl.style = "Table Grid"
        for ci, hdr in enumerate(["Priority", "ID", "Title", "Severity", "Effort/Owner"]):
            c = rr_tbl.cell(0, ci)
            _set_cell_shading(c, "C62828")
            p = c.paragraphs[0]
            r = p.add_run(hdr)
            r.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)
        for rank, f in enumerate(sorted_findings, 1):
            sev = (f.get("severity") or "").lower()
            pri = priority_labels.get(sev, "P5")
            row_cells = [rr_tbl.cell(rank, ci) for ci in range(5)]
            row_cells[0].paragraphs[0].add_run(pri)
            row_cells[1].paragraphs[0].add_run(_safe(f.get("finding_id")))
            row_cells[2].paragraphs[0].add_run(_safe(f.get("title")))
            sc = row_cells[3]
            _set_cell_shading(sc, SEV_DOCX_COLORS.get(sev, "757575"))
            sp = sc.paragraphs[0]
            sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sr = sp.add_run(sev.capitalize())
            sr.font.color.rgb = RGBColor(255, 255, 255)
            sr.bold = True
            eo = f"{effort_map.get(sev, 'M')} | {owner_map.get(sev, 'Dev Team')}"
            row_cells[4].paragraphs[0].add_run(eo)
        doc.add_paragraph()

    # Section 3: Detailed Remediation Steps
    _add_section_header("3. Detailed Remediation Steps")
    for fi, f in enumerate(sorted_findings):
        sev = (f.get("severity") or "").lower()
        fid = _safe(f.get("finding_id"), f"F-{fi + 1:02d}")
        rs = (f.get("retest_status") or "pending").lower()

        f_tbl = doc.add_table(rows=1, cols=1)
        f_tbl.style = "Table Grid"
        hdr_cell = f_tbl.cell(0, 0)
        _set_cell_shading(hdr_cell, SEV_DOCX_COLORS.get(sev, "757575"))
        hp = hdr_cell.paragraphs[0]
        hr3 = hp.add_run(f"{fid}  —  {_safe(f.get('title'))}")
        hr3.bold = True
        hr3.font.size = Pt(11)
        hr3.font.color.rgb = RGBColor(255, 255, 255)

        if f.get("description"):
            _add_para("Context:", bold=True, color="C62828")
            _add_para(_safe(f.get("description"))[:400])
        if f.get("impact"):
            _add_para("Business Impact:", bold=True, color="C62828")
            _add_para(_safe(f.get("impact"))[:300])

        rec = _safe(f.get("recommendation"))
        if rec and rec != "—":
            _add_para("Immediate Actions:", bold=True, color="C62828")
            for step_i, line in enumerate(rec.split("\n"), 1):
                if line.strip():
                    p = doc.add_paragraph(style="List Number")
                    p.add_run(line.strip())

        _add_para("Verification Steps:", bold=True, color="C62828")
        for step in [
            "Apply the recommended fix in a test environment.",
            "Re-run the relevant test case or scan.",
            "Document evidence of the fix (screenshots, config exports).",
            "Submit for retest validation by the security team.",
        ]:
            p = doc.add_paragraph(style="List Number")
            p.add_run(step)

        _add_para(f"Effort: {effort_map.get(sev, 'Medium')}  |  Priority: {priority_labels.get(sev, 'P3')}",
                  color="546E7A")
        doc.add_paragraph()

    # Section 4: Verification Checklist
    _add_section_header("4. Verification Checklist")
    if sorted_findings:
        chk_tbl = doc.add_table(rows=len(sorted_findings) + 1, cols=5)
        chk_tbl.style = "Table Grid"
        for ci, hdr in enumerate(["ID", "Title", "Severity", "Fix Applied", "Verified/Closed"]):
            c = chk_tbl.cell(0, ci)
            _set_cell_shading(c, "C62828")
            p = c.paragraphs[0]
            r = p.add_run(hdr)
            r.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)
        for i, f in enumerate(sorted_findings):
            sev = (f.get("severity") or "").lower()
            row_cells = [chk_tbl.cell(i + 1, ci) for ci in range(5)]
            row_cells[0].paragraphs[0].add_run(_safe(f.get("finding_id")))
            row_cells[1].paragraphs[0].add_run(_safe(f.get("title")))
            sc = row_cells[2]
            _set_cell_shading(sc, SEV_DOCX_COLORS.get(sev, "757575"))
            sp = sc.paragraphs[0]
            sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sr = sp.add_run(sev.capitalize())
            sr.font.color.rgb = RGBColor(255, 255, 255)
            sr.bold = True
            for ci2 in [3, 4]:
                cp = row_cells[ci2].paragraphs[0]
                cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cp.add_run("☐")
        doc.add_paragraph()

    # Section 5: Sign-off
    _add_section_header("5. Sign-off")
    sign_tbl = doc.add_table(rows=4, cols=4)
    sign_tbl.style = "Table Grid"
    for ci, hdr in enumerate(["Role", "Name", "Signature", "Date"]):
        c = sign_tbl.cell(0, ci)
        _set_cell_shading(c, "C62828")
        p = c.paragraphs[0]
        r = p.add_run(hdr)
        r.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
    for i, (role, name) in enumerate([
        ("Prepared By", _safe(report.get("prepared_by"))),
        ("Reviewed By", _safe(report.get("reviewed_by"))),
        ("Client Sign-off", ""),
    ], 1):
        sign_tbl.cell(i, 0).paragraphs[0].add_run(role)
        sign_tbl.cell(i, 1).paragraphs[0].add_run(name)
        sign_tbl.cell(i, 2).paragraphs[0].add_run("")
        sign_tbl.cell(i, 3).paragraphs[0].add_run("")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
