"""Risk Register Export Service — PDF and DOCX."""
import io
import json
from datetime import datetime
from typing import Any, List, Optional

from reportlab.lib import colors
from reportlab.lib.colors import HexColor
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import (
    BaseDocTemplate, Frame, HRFlowable, KeepTogether, PageBreak,
    PageTemplate, Paragraph, Spacer, Table, TableStyle,
)

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ── Colours ───────────────────────────────────────────────────────────────────
NAVY = HexColor("#1A237E")
BLUE = HexColor("#1565C0")
DARK_TEXT = HexColor("#37474F")
WHITE = colors.white
GREY_BG = HexColor("#ECEFF1")

LEVEL_COLORS = {
    "critical":     HexColor("#C62828"),
    "high":         HexColor("#E64A19"),
    "medium_high":  HexColor("#EF6C00"),
    "medium":       HexColor("#F9A825"),
    "low":          HexColor("#2E7D32"),
}
LEVEL_COLORS_HEX = {
    "critical": "#C62828", "high": "#E64A19",
    "medium_high": "#EF6C00", "medium": "#F9A825", "low": "#2E7D32",
}
MEASURE_STATUS_COLORS = {
    "in_place":    HexColor("#2E7D32"),
    "not_possible": HexColor("#C62828"),
    "pending":     HexColor("#757575"),
}
MEASURE_STATUS_HEX = {
    "in_place": "#2E7D32", "not_possible": "#C62828", "pending": "#757575",
}
MEASURE_STATUS_LABEL = {
    "in_place": "✓ In Place", "not_possible": "✗ Not Possible", "pending": "○ Pending",
}

FACTOR_NAMES = [
    ("accessibility", "Accessibility"),
    ("discoverability", "Discoverability"),
    ("exploitability", "Exploitability"),
    ("authentication_score", "Authentication"),
    ("authentication", "Authentication"),
    ("repeatability", "Repeatability"),
]


def _safe(val: Any, fallback: str = "—") -> str:
    if val is None or (isinstance(val, str) and not val.strip()):
        return fallback
    return str(val)


def _level_str(risk: Any) -> str:
    lv = getattr(risk, "residual_risk_level", None) or ""
    if not lv:
        score = getattr(risk, "risk_matrix_score", None) or 0
        if score >= 21: lv = "critical"
        elif score >= 13: lv = "high"
        elif score >= 10: lv = "medium_high"
        elif score >= 5: lv = "medium"
        else: lv = "low"
    return lv.lower()


def _parse_json(val: Any) -> Any:
    if not val:
        return None
    if isinstance(val, (dict, list)):
        return val
    try:
        return json.loads(val)
    except Exception:
        return None


def _fmt_date(val: Any) -> str:
    if not val:
        return "—"
    if isinstance(val, datetime):
        return val.strftime("%d %b %Y")
    try:
        return datetime.fromisoformat(str(val)).strftime("%d %b %Y")
    except Exception:
        return str(val)


# ── PDF helpers ───────────────────────────────────────────────────────────────

def _make_styles():
    base = getSampleStyleSheet()
    normal = base["Normal"]
    styles = {
        "title": ParagraphStyle("title", parent=normal, fontSize=22, fontName="Helvetica-Bold",
                                textColor=NAVY, spaceAfter=6),
        "subtitle": ParagraphStyle("subtitle", parent=normal, fontSize=11, fontName="Helvetica",
                                   textColor=DARK_TEXT, spaceAfter=12),
        "h2": ParagraphStyle("h2", parent=normal, fontSize=14, fontName="Helvetica-Bold",
                              textColor=NAVY, spaceBefore=14, spaceAfter=6),
        "h3": ParagraphStyle("h3", parent=normal, fontSize=11, fontName="Helvetica-Bold",
                              textColor=DARK_TEXT, spaceBefore=8, spaceAfter=4),
        "body": ParagraphStyle("body", parent=normal, fontSize=9, fontName="Helvetica",
                               textColor=DARK_TEXT, spaceAfter=4, leading=13),
        "caption": ParagraphStyle("caption", parent=normal, fontSize=8, fontName="Helvetica",
                                  textColor=HexColor("#78909C"), spaceAfter=2),
        "label": ParagraphStyle("label", parent=normal, fontSize=8, fontName="Helvetica-Bold",
                                textColor=DARK_TEXT, spaceAfter=2),
    }
    return styles


class _PageTemplate:
    def __init__(self, client_name: str):
        self.client_name = client_name

    def on_page(self, canvas, doc):
        canvas.saveState()
        canvas.setFillColor(NAVY)
        canvas.rect(0, A4[1] - 1.2 * cm, A4[0], 1.2 * cm, fill=1, stroke=0)
        canvas.setFillColor(WHITE)
        canvas.setFont("Helvetica-Bold", 9)
        canvas.drawString(1.5 * cm, A4[1] - 0.8 * cm, f"Risk Register — {self.client_name}")
        canvas.setFont("Helvetica", 8)
        canvas.drawRightString(A4[0] - 1.5 * cm, A4[1] - 0.8 * cm,
                               f"Generated {datetime.now().strftime('%d %b %Y')}")
        canvas.setFillColor(GREY_BG)
        canvas.rect(0, 0, A4[0], 0.8 * cm, fill=1, stroke=0)
        canvas.setFillColor(DARK_TEXT)
        canvas.setFont("Helvetica", 8)
        canvas.drawCentredString(A4[0] / 2, 0.3 * cm, f"Page {doc.page}")
        canvas.restoreState()


def _build_pdf_risk_detail(risk: Any, styles: dict) -> list:
    """Build PDF flowable blocks for a single risk."""
    lv = _level_str(risk)
    lv_color = LEVEL_COLORS.get(lv, HexColor("#888888"))
    score = getattr(risk, "risk_matrix_score", None) or 0
    treatment = _safe(getattr(risk, "treatment_option", None)).capitalize()
    wizard = _parse_json(getattr(risk, "wizard_data_json", None)) or {}
    measures = _parse_json(getattr(risk, "measures_json", None)) or []
    ai_data = _parse_json(getattr(risk, "ai_assessment_json", None)) or {}
    workarounds = {w["measure_id"]: w["alternative"] for w in ai_data.get("workarounds", [])}
    lf = ai_data.get("likelihood_factors", {})

    elems = []
    # Risk header bar
    header_data = [[
        Paragraph(f'<font color="#FFFFFF"><b>{lv.upper().replace("_", "-")}</b></font>', styles["label"]),
        Paragraph(f'<font color="#FFFFFF"><b>{risk.title}</b></font>', styles["body"]),
        Paragraph(f'<font color="#FFFFFF">Score: {score}/25 · {treatment}</font>', styles["caption"]),
    ]]
    header_table = Table(header_data, colWidths=[2.5 * cm, None, 3.5 * cm])
    header_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), lv_color),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ("LEFTPADDING", (0, 0), (0, 0), 8),
    ]))
    elems.append(header_table)
    elems.append(Spacer(1, 4))

    # Description / scenario
    if risk.description:
        elems.append(Paragraph(f"<b>Description:</b> {risk.description}", styles["body"]))

    # Likelihood factors table
    factor_rows = [["Factor", "Score", "Rationale"]]
    factor_keys = [
        ("accessibility", "Accessibility"),
        ("discoverability", "Discoverability"),
        ("exploitability", "Exploitability"),
        ("authentication", "Authentication"),
        ("repeatability", "Repeatability"),
    ]
    for key, label in factor_keys:
        score_val = lf.get(key) or getattr(risk, key if key != "authentication" else "authentication_score", None) or 3
        rationale = lf.get(f"{key}_rationale", "—")
        factor_rows.append([label, f"{score_val}/5", Paragraph(rationale, styles["caption"])])

    consequence_val = getattr(risk, "consequence", None) or 3
    factor_rows.append(["Consequence", f"{consequence_val}/5",
                         Paragraph(lf.get("consequence_rationale", ai_data.get("consequence_rationale", "—")), styles["caption"])])

    ft = Table(factor_rows, colWidths=[3.5 * cm, 1.5 * cm, None])
    ft.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), NAVY),
        ("TEXTCOLOR", (0, 0), (-1, 0), WHITE),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [WHITE, GREY_BG]),
        ("GRID", (0, 0), (-1, -1), 0.3, HexColor("#CFD8DC")),
        ("TOPPADDING", (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
        ("LEFTPADDING", (0, 0), (-1, -1), 5),
    ]))
    elems.append(Spacer(1, 4))
    elems.append(Paragraph("Likelihood Assessment", styles["h3"]))
    elems.append(ft)
    elems.append(Spacer(1, 4))

    # Measures checklist
    if measures:
        elems.append(Paragraph("Security Measures Checklist", styles["h3"]))
        m_rows = [["#", "Measure", "Status"]]
        for m in measures:
            status = m.get("status", "pending")
            status_label = MEASURE_STATUS_LABEL.get(status, status)
            m_rows.append([
                m.get("id", ""),
                Paragraph(m.get("text", ""), styles["body"]),
                Paragraph(f'<font color="{MEASURE_STATUS_HEX.get(status, "#555")}">{status_label}</font>', styles["body"]),
            ])
            # Workaround if not_possible
            wa = workarounds.get(m.get("id", ""))
            if wa and status == "not_possible":
                m_rows.append(["", Paragraph(f'<i>Workaround: {wa}</i>', styles["caption"]), ""])

        mt = Table(m_rows, colWidths=[1 * cm, None, 3 * cm])
        mt.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), BLUE),
            ("TEXTCOLOR", (0, 0), (-1, 0), WHITE),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 8),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [WHITE, GREY_BG]),
            ("GRID", (0, 0), (-1, -1), 0.3, HexColor("#CFD8DC")),
            ("TOPPADDING", (0, 0), (-1, -1), 3),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ("LEFTPADDING", (0, 0), (-1, -1), 5),
        ]))
        elems.append(mt)
        elems.append(Spacer(1, 4))

    # AI commentary
    commentary = ai_data.get("overall_commentary", "")
    if commentary:
        elems.append(Paragraph("AI Commentary", styles["h3"]))
        elems.append(Paragraph(commentary, styles["body"]))

    elems.append(HRFlowable(width="100%", thickness=0.5, color=HexColor("#CFD8DC"), spaceAfter=8))
    return elems


def generate_risk_register_pdf(client_name: str, risks: list) -> io.BytesIO:
    buf = io.BytesIO()
    doc = BaseDocTemplate(buf, pagesize=A4,
                          leftMargin=1.5 * cm, rightMargin=1.5 * cm,
                          topMargin=2 * cm, bottomMargin=1.5 * cm)
    pt = _PageTemplate(client_name)
    frame = Frame(doc.leftMargin, doc.bottomMargin,
                  doc.width, doc.height, id="main")
    doc.addPageTemplates([PageTemplate(id="main", frames=[frame], onPage=pt.on_page)])

    styles = _make_styles()
    story = []

    # Cover
    story.append(Spacer(1, 1 * cm))
    story.append(Paragraph("Risk Register", styles["title"]))
    story.append(Paragraph(f"Client: {client_name}", styles["subtitle"]))
    story.append(Paragraph(f"Generated: {datetime.now().strftime('%d %B %Y %H:%M')}", styles["subtitle"]))
    story.append(Paragraph(f"Total risks: {len(risks)}", styles["caption"]))
    story.append(HRFlowable(width="100%", thickness=1, color=NAVY, spaceAfter=12))

    # Summary table
    story.append(Paragraph("Risk Summary", styles["h2"]))
    summary_rows = [["Level", "Title", "Score", "Treatment", "Status"]]
    for r in risks:
        lv = _level_str(r)
        score = getattr(r, "risk_matrix_score", None) or 0
        treatment = _safe(getattr(r, "treatment_option", None)).capitalize()
        summary_rows.append([
            Paragraph(f'<font color="{LEVEL_COLORS_HEX.get(lv, "#888")}">'
                      f'<b>{lv.upper().replace("_", "-")}</b></font>', styles["body"]),
            Paragraph(r.title[:80], styles["body"]),
            str(score),
            treatment or "—",
            _safe(r.status).capitalize(),
        ])

    st = Table(summary_rows, colWidths=[2.5 * cm, None, 1.5 * cm, 2.5 * cm, 2 * cm])
    st.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), NAVY),
        ("TEXTCOLOR", (0, 0), (-1, 0), WHITE),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [WHITE, GREY_BG]),
        ("GRID", (0, 0), (-1, -1), 0.3, HexColor("#CFD8DC")),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("LEFTPADDING", (0, 0), (-1, -1), 5),
    ]))
    story.append(st)
    story.append(PageBreak())

    # Detail sections
    story.append(Paragraph("Risk Details", styles["h2"]))
    for risk in risks:
        elems = _build_pdf_risk_detail(risk, styles)
        story.append(KeepTogether(elems[:4]))
        story.extend(elems[4:])

    doc.build(story)
    buf.seek(0)
    return buf


def generate_single_risk_pdf(client_name: str, risk: Any) -> io.BytesIO:
    buf = io.BytesIO()
    doc = BaseDocTemplate(buf, pagesize=A4,
                          leftMargin=1.5 * cm, rightMargin=1.5 * cm,
                          topMargin=2 * cm, bottomMargin=1.5 * cm)
    pt = _PageTemplate(client_name)
    frame = Frame(doc.leftMargin, doc.bottomMargin, doc.width, doc.height, id="main")
    doc.addPageTemplates([PageTemplate(id="main", frames=[frame], onPage=pt.on_page)])

    styles = _make_styles()
    story = [Spacer(1, 0.5 * cm)]
    story.append(Paragraph("Risk Detail Report", styles["title"]))
    story.append(Paragraph(f"Client: {client_name} · {datetime.now().strftime('%d %B %Y')}", styles["subtitle"]))
    story.append(HRFlowable(width="100%", thickness=1, color=NAVY, spaceAfter=12))
    story.extend(_build_pdf_risk_detail(risk, styles))

    doc.build(story)
    buf.seek(0)
    return buf


# ── DOCX helpers ──────────────────────────────────────────────────────────────

def _docx_add_colored_heading(doc: Document, text: str, level: int, hex_color: str = "1A237E"):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor.from_string(hex_color)


def _docx_risk_detail(doc: Document, risk: Any):
    lv = _level_str(risk)
    lv_hex = LEVEL_COLORS_HEX.get(lv, "888888").lstrip("#")
    score = getattr(risk, "risk_matrix_score", None) or 0
    treatment = _safe(getattr(risk, "treatment_option", None)).capitalize()
    wizard = _parse_json(getattr(risk, "wizard_data_json", None)) or {}
    measures = _parse_json(getattr(risk, "measures_json", None)) or []
    ai_data = _parse_json(getattr(risk, "ai_assessment_json", None)) or {}
    workarounds = {w["measure_id"]: w["alternative"] for w in ai_data.get("workarounds", [])}
    lf = ai_data.get("likelihood_factors", {})

    h = doc.add_heading(risk.title, level=2)
    for run in h.runs:
        run.font.color.rgb = RGBColor.from_string(lv_hex)

    p = doc.add_paragraph()
    p.add_run(f"Level: {lv.upper().replace('_', '-')}  |  Score: {score}/25  |  Treatment: {treatment or '—'}").bold = True

    if risk.description:
        doc.add_paragraph(risk.description)

    # Likelihood factors table
    doc.add_heading("Likelihood Assessment", level=3)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Factor"
    hdr[1].text = "Score"
    hdr[2].text = "Rationale"
    for cell in hdr:
        for run in cell.paragraphs[0].runs:
            run.bold = True

    factor_keys = [
        ("accessibility", "Accessibility"),
        ("discoverability", "Discoverability"),
        ("exploitability", "Exploitability"),
        ("authentication", "Authentication"),
        ("repeatability", "Repeatability"),
        ("consequence", "Consequence"),
    ]
    for key, label in factor_keys:
        attr_key = key if key != "authentication" else "authentication_score"
        val = lf.get(key) or getattr(risk, attr_key, None) or 3
        if key == "consequence":
            val = getattr(risk, "consequence", None) or 3
        rationale = lf.get(f"{key}_rationale", ai_data.get("consequence_rationale", "—") if key == "consequence" else "—")
        row = table.add_row().cells
        row[0].text = label
        row[1].text = f"{val}/5"
        row[2].text = rationale

    # Measures checklist
    if measures:
        doc.add_heading("Security Measures", level=3)
        for m in measures:
            status = m.get("status", "pending")
            label = MEASURE_STATUS_LABEL.get(status, status)
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f"{label}  ").bold = True
            p.add_run(m.get("text", ""))
            wa = workarounds.get(m.get("id", ""))
            if wa and status == "not_possible":
                wp = doc.add_paragraph(style="List Bullet 2")
                wp.add_run(f"Workaround: {wa}").italic = True

    # AI Commentary
    commentary = ai_data.get("overall_commentary", "")
    if commentary:
        doc.add_heading("AI Commentary", level=3)
        doc.add_paragraph(commentary)

    doc.add_paragraph("─" * 60)


def generate_risk_register_docx(client_name: str, risks: list) -> io.BytesIO:
    doc = Document()
    doc.add_heading("Risk Register", 0)
    p = doc.add_paragraph(f"Client: {client_name}  |  Generated: {datetime.now().strftime('%d %B %Y %H:%M')}")

    # Summary table
    doc.add_heading("Summary", level=1)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, col in enumerate(["Level", "Title", "Score", "Treatment", "Status"]):
        hdr[i].text = col
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
    for r in risks:
        lv = _level_str(r)
        score = getattr(r, "risk_matrix_score", None) or 0
        treatment = _safe(getattr(r, "treatment_option", None)).capitalize()
        row = table.add_row().cells
        row[0].text = lv.upper().replace("_", "-")
        row[1].text = r.title[:80]
        row[2].text = str(score)
        row[3].text = treatment or "—"
        row[4].text = _safe(r.status).capitalize()

    doc.add_page_break()
    doc.add_heading("Risk Details", level=1)
    for risk in risks:
        _docx_risk_detail(doc, risk)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def generate_single_risk_docx(client_name: str, risk: Any) -> io.BytesIO:
    doc = Document()
    doc.add_heading("Risk Detail Report", 0)
    doc.add_paragraph(f"Client: {client_name}  |  {datetime.now().strftime('%d %B %Y')}")
    _docx_risk_detail(doc, risk)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
