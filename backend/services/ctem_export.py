"""CTEM report export — PDF (reportlab) and DOCX (python-docx)."""
import io
from datetime import datetime, timezone
from typing import Dict, Optional, List

PHASES = ["scope", "discover", "prioritise", "validate", "mobilise"]
PHASE_LABELS = {
    "scope": "Phase 1: Scope",
    "discover": "Phase 2: Discover",
    "prioritise": "Phase 3: Prioritise",
    "validate": "Phase 4: Validate",
    "mobilise": "Phase 5: Mobilise",
}


def _fmt_dt(dt) -> str:
    if not dt:
        return "—"
    if isinstance(dt, str):
        return dt[:10]
    return dt.strftime("%Y-%m-%d")


# ── PDF ───────────────────────────────────────────────────────────────────────

def generate_ctem_pdf(program, phase_map: Dict) -> io.BytesIO:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.lib import colors
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable,
    )

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=20*mm, rightMargin=20*mm, topMargin=20*mm, bottomMargin=20*mm,
    )
    styles = getSampleStyleSheet()
    BRAND = colors.HexColor("#4285F4")
    DARK = colors.HexColor("#1a1a2e")
    PHASE_COLOR = colors.HexColor("#E8F0FE")

    title_style = ParagraphStyle("Title", parent=styles["Title"], textColor=BRAND, fontSize=20, spaceAfter=4)
    h1_style = ParagraphStyle("H1", parent=styles["Heading1"], textColor=BRAND, fontSize=14, spaceBefore=14, spaceAfter=4)
    h2_style = ParagraphStyle("H2", parent=styles["Heading2"], textColor=DARK, fontSize=11, spaceBefore=8, spaceAfter=2)
    body_style = ParagraphStyle("Body", parent=styles["Normal"], fontSize=9, leading=13, spaceAfter=4)
    label_style = ParagraphStyle("Label", parent=styles["Normal"], fontSize=8, textColor=colors.grey)

    story = []

    # Cover
    story.append(Paragraph("Continuous Threat Exposure Management", title_style))
    story.append(Paragraph(f"<b>{program.name}</b>", h1_style))
    if program.description:
        story.append(Paragraph(program.description, body_style))
    meta_data = [
        ["Status", program.status or "active", "Created by", program.created_by or "—"],
        ["Created", _fmt_dt(program.created_at), "Generated", datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")],
        ["Current Phase", program.current_phase or "scope", "", ""],
    ]
    meta_table = Table(meta_data, colWidths=[35*mm, 50*mm, 35*mm, 50*mm])
    meta_table.setStyle(TableStyle([
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("TEXTCOLOR", (0, 0), (0, -1), colors.grey),
        ("TEXTCOLOR", (2, 0), (2, -1), colors.grey),
        ("FONTNAME", (1, 0), (1, -1), "Helvetica-Bold"),
        ("FONTNAME", (3, 0), (3, -1), "Helvetica-Bold"),
        ("GRID", (0, 0), (-1, -1), 0.25, colors.lightgrey),
        ("BACKGROUND", (0, 0), (-1, -1), colors.whitesmoke),
        ("PADDING", (0, 0), (-1, -1), 4),
    ]))
    story.append(meta_table)
    story.append(Spacer(1, 8*mm))
    story.append(HRFlowable(width="100%", thickness=1, color=BRAND))

    def tbl(data, col_widths=None, header_bg=PHASE_COLOR):
        if not data:
            return None
        t = Table(data, colWidths=col_widths)
        style = [
            ("BACKGROUND", (0, 0), (-1, 0), header_bg),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 8),
            ("GRID", (0, 0), (-1, -1), 0.25, colors.lightgrey),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F8F9FA")]),
            ("PADDING", (0, 0), (-1, -1), 4),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ]
        t.setStyle(TableStyle(style))
        return t

    # Phases
    for phase_key in PHASES:
        pn = phase_map.get(phase_key)
        pd = (pn.phase_data_json or {}) if pn else {}
        story.append(Spacer(1, 6*mm))
        story.append(Paragraph(PHASE_LABELS[phase_key], h1_style))
        done_text = f"✓ Completed by {pn.completed_by} on {_fmt_dt(pn.completed_at)}" if (pn and pn.completed) else "In Progress"
        story.append(Paragraph(done_text, label_style))
        story.append(Spacer(1, 2*mm))

        if phase_key == "scope":
            assets: List[Dict] = pd.get("assets", [])
            if assets:
                story.append(Paragraph("Scoped Assets", h2_style))
                rows = [["Resource", "Type", "Status", "Notes"]]
                for a in assets:
                    rows.append([
                        Paragraph(str(a.get("resource_id", "")), body_style),
                        Paragraph(str(a.get("resource_type", "")), body_style),
                        str(a.get("scope_status", "untagged")).replace("_", " ").title(),
                        Paragraph(str(a.get("notes", "")), body_style),
                    ])
                t = tbl(rows, col_widths=[60*mm, 35*mm, 30*mm, 45*mm])
                if t:
                    story.append(t)

        elif phase_key == "discover":
            cats: List[Dict] = pd.get("categories", [])
            if cats:
                story.append(Paragraph("Exposure Categories", h2_style))
                rows = [["Category", "Total", "Critical", "High", "Medium", "Low"]]
                for c in cats:
                    rows.append([c.get("category",""), c.get("total",0), c.get("critical",0), c.get("high",0), c.get("medium",0), c.get("low",0)])
                t = tbl(rows, col_widths=[65*mm, 18*mm, 20*mm, 18*mm, 20*mm, 18*mm])
                if t:
                    story.append(t)
            if pd.get("summary"):
                story.append(Paragraph("Summary", h2_style))
                story.append(Paragraph(pd["summary"], body_style))

        elif phase_key == "prioritise":
            items: List[Dict] = pd.get("items", [])
            if items:
                story.append(Paragraph("Priority Items", h2_style))
                rows = [["#", "Title", "Severity", "Source", "Rationale", "Analyst Notes"]]
                for item in items:
                    rows.append([
                        str(item.get("rank", "")),
                        Paragraph(str(item.get("title", "")), body_style),
                        str(item.get("severity", "")),
                        str(item.get("source", "")),
                        Paragraph(str(item.get("rationale", "")), body_style),
                        Paragraph(str(item.get("analyst_notes", "")), body_style),
                    ])
                t = tbl(rows, col_widths=[8*mm, 50*mm, 20*mm, 18*mm, 40*mm, 34*mm])
                if t:
                    story.append(t)

        elif phase_key == "validate":
            methods: List[Dict] = pd.get("methods", [])
            if methods:
                story.append(Paragraph("Validation Methods", h2_style))
                rows = [["Method", "Tests Run", "Confirmed Exploitable", "Notes"]]
                for m in methods:
                    rows.append([
                        Paragraph(str(m.get("name", "")), body_style),
                        str(m.get("tests_run", "")),
                        str(m.get("confirmed", "")),
                        Paragraph(str(m.get("notes", "")), body_style),
                    ])
                t = tbl(rows, col_widths=[60*mm, 28*mm, 40*mm, 42*mm])
                if t:
                    story.append(t)
            if pd.get("notable_findings"):
                story.append(Paragraph("Notable Findings", h2_style))
                for line in str(pd["notable_findings"]).split("\n"):
                    if line.strip():
                        story.append(Paragraph(f"• {line.strip().lstrip('•-').strip()}", body_style))

        elif phase_key == "mobilise":
            owners: List[Dict] = pd.get("owners", [])
            if owners:
                story.append(Paragraph("Owner Teams", h2_style))
                rows = [["Team", "Open", "Closed On-Time", "SLA Breach", "Notes"]]
                for o in owners:
                    rows.append([
                        Paragraph(str(o.get("team", "")), body_style),
                        str(o.get("open", "")),
                        str(o.get("closed_on_time", "")),
                        str(o.get("sla_breach", "")),
                        Paragraph(str(o.get("notes", "")), body_style),
                    ])
                t = tbl(rows, col_widths=[50*mm, 20*mm, 30*mm, 22*mm, 48*mm])
                if t:
                    story.append(t)
            blockers: List[str] = pd.get("blockers", [])
            if blockers:
                story.append(Paragraph("Blockers", h2_style))
                for b in blockers:
                    if b.strip():
                        story.append(Paragraph(f"• {b.strip()}", body_style))

        # AI brief
        if pn and pn.ai_brief:
            story.append(Paragraph("AI Analysis", h2_style))
            for para in pn.ai_brief.split("\n\n"):
                if para.strip():
                    story.append(Paragraph(para.strip(), body_style))

        # Analyst notes
        if pn and pn.notes:
            story.append(Paragraph("Analyst Notes", h2_style))
            story.append(Paragraph(pn.notes, body_style))

    doc.build(story)
    buf.seek(0)
    return buf


# ── DOCX ──────────────────────────────────────────────────────────────────────

def generate_ctem_docx(program, phase_map: Dict) -> io.BytesIO:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    BRAND = RGBColor(0x42, 0x85, 0xF4)

    def _heading(text: str, level: int = 1):
        h = doc.add_heading(text, level=level)
        h.runs[0].font.color.rgb = BRAND if level == 1 else RGBColor(0x1a, 0x1a, 0x2e)

    def _add_table(headers: List[str], rows: List[List[str]]):
        t = doc.add_table(rows=1 + len(rows), cols=len(headers))
        t.style = "Table Grid"
        hdr = t.rows[0].cells
        for i, h in enumerate(headers):
            hdr[i].text = h
            hdr[i].paragraphs[0].runs[0].bold = True
        for ri, row in enumerate(rows):
            for ci, val in enumerate(row):
                t.rows[ri + 1].cells[ci].text = str(val)

    # Cover
    title = doc.add_heading("Continuous Threat Exposure Management", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading(program.name, 1)
    if program.description:
        doc.add_paragraph(program.description)
    _add_table(
        ["Field", "Value"],
        [
            ["Status", program.status or "active"],
            ["Created by", program.created_by or "—"],
            ["Created", _fmt_dt(program.created_at)],
            ["Current Phase", program.current_phase or "scope"],
            ["Generated", datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")],
        ],
    )
    doc.add_page_break()

    for phase_key in PHASES:
        pn = phase_map.get(phase_key)
        pd = (pn.phase_data_json or {}) if pn else {}
        _heading(PHASE_LABELS[phase_key], 1)
        done_text = f"Completed by {pn.completed_by} on {_fmt_dt(pn.completed_at)}" if (pn and pn.completed) else "Status: In Progress"
        doc.add_paragraph(done_text).italic = True

        if phase_key == "scope":
            assets = pd.get("assets", [])
            if assets:
                _heading("Scoped Assets", 2)
                _add_table(
                    ["Resource", "Type", "Status", "Notes"],
                    [[a.get("resource_id",""), a.get("resource_type",""), a.get("scope_status","").replace("_"," ").title(), a.get("notes","")] for a in assets],
                )

        elif phase_key == "discover":
            cats = pd.get("categories", [])
            if cats:
                _heading("Exposure Categories", 2)
                _add_table(
                    ["Category", "Total", "Critical", "High", "Medium", "Low"],
                    [[c.get("category",""), c.get("total",0), c.get("critical",0), c.get("high",0), c.get("medium",0), c.get("low",0)] for c in cats],
                )

        elif phase_key == "prioritise":
            items = pd.get("items", [])
            if items:
                _heading("Priority Items", 2)
                _add_table(
                    ["#", "Title", "Severity", "Source", "Rationale", "Analyst Notes"],
                    [[item.get("rank",""), item.get("title",""), item.get("severity",""), item.get("source",""), item.get("rationale",""), item.get("analyst_notes","")] for item in items],
                )

        elif phase_key == "validate":
            methods = pd.get("methods", [])
            if methods:
                _heading("Validation Methods", 2)
                _add_table(
                    ["Method", "Tests Run", "Confirmed Exploitable", "Notes"],
                    [[m.get("name",""), m.get("tests_run",""), m.get("confirmed",""), m.get("notes","")] for m in methods],
                )
            if pd.get("notable_findings"):
                _heading("Notable Findings", 2)
                for line in str(pd["notable_findings"]).split("\n"):
                    if line.strip():
                        doc.add_paragraph(line.strip(), style="List Bullet")

        elif phase_key == "mobilise":
            owners = pd.get("owners", [])
            if owners:
                _heading("Owner Teams", 2)
                _add_table(
                    ["Team", "Open", "Closed On-Time", "SLA Breach", "Notes"],
                    [[o.get("team",""), o.get("open",""), o.get("closed_on_time",""), o.get("sla_breach",""), o.get("notes","")] for o in owners],
                )
            blockers = pd.get("blockers", [])
            if blockers:
                _heading("Blockers", 2)
                for b in blockers:
                    if b.strip():
                        doc.add_paragraph(b.strip(), style="List Bullet")

        if pn and pn.ai_brief:
            _heading("AI Analysis", 2)
            doc.add_paragraph(pn.ai_brief)

        if pn and pn.notes:
            _heading("Analyst Notes", 2)
            doc.add_paragraph(pn.notes)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
